from __future__ import annotations

import base64
from io import BytesIO

from docx import Document
from PIL import Image as PILImage
from pypdf import PdfReader

from app.models.project import (
    ExportFormat,
    FullExamTranslationAcceptanceStatus,
    FullExamTranslationReport,
    MarksPolicy,
    ProjectMetadata,
    ProjectSession,
    QuestionAssetInfo,
    QuestionItem,
    QuestionStatus,
)
from app.services.export import (
    build_project_docx_bytes,
    build_project_pdf_bytes,
)
from app.services.full_exam_export import build_full_exam_export_report
from app.services.readiness import build_project_readiness_report


def _portrait_asset() -> QuestionAssetInfo:
    image = PILImage.new("RGB", (400, 1200), (255, 255, 255))
    output = BytesIO()
    image.save(output, format="PNG")
    payload = output.getvalue()
    return QuestionAssetInfo(
        id="asset-1",
        name="portrait.png",
        size=len(payload),
        type="image/png",
        data_base64=base64.b64encode(payload).decode("ascii"),
    )


def _translation_report(
    status: FullExamTranslationAcceptanceStatus,
) -> FullExamTranslationReport:
    return FullExamTranslationReport(
        status=status,
        total_questions=1,
        active_questions=1,
        translated_questions=1,
        accepted_questions=(
            1
            if status == FullExamTranslationAcceptanceStatus.accepted
            else 0
        ),
        needs_review_questions=(
            1
            if status == FullExamTranslationAcceptanceStatus.needs_review
            else 0
        ),
        completion_percent=100,
        total_items=1,
        translated_items=1,
    )


def _project(
    *,
    policy: MarksPolicy,
    translation_status: FullExamTranslationAcceptanceStatus,
    with_asset: bool = True,
) -> ProjectSession:
    return ProjectSession(
        metadata=ProjectMetadata(
            paper_title="Marks policy",
            total_marks="20",
            marks_policy=policy,
            export_formats=[ExportFormat.docx, ExportFormat.pdf],
        ),
        questions=[
            QuestionItem(
                id="question-1",
                original_number="1",
                original_text="Use Fig. 1.1. [80]",
                translated_text="استخدم الشكل 1.1. [80]",
                marks=80,
                status=QuestionStatus.approved,
                order_index=1,
                attachments=[_portrait_asset()] if with_asset else [],
            )
        ],
        full_exam_translation_report=_translation_report(
            translation_status
        ),
    )


def _docx_text(payload: bytes) -> str:
    document = Document(BytesIO(payload))
    values: list[str] = []
    for paragraph in document.paragraphs:
        values.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def test_unresolved_marks_policy_keeps_readiness_mismatch() -> None:
    project = _project(
        policy=MarksPolicy.unresolved,
        translation_status=FullExamTranslationAcceptanceStatus.accepted,
    )

    report = build_project_readiness_report(project)

    assert any(
        issue.code == "paper_total_marks_mismatch"
        for issue in report.issues
    )


def test_question_total_policy_resolves_mismatch_and_exports_80() -> None:
    project = _project(
        policy=MarksPolicy.use_question_total,
        translation_status=FullExamTranslationAcceptanceStatus.accepted,
    )

    readiness = build_project_readiness_report(project)
    docx = build_project_docx_bytes(project)
    report = build_full_exam_export_report(
        project,
        ExportFormat.docx,
        docx,
    )
    marks_check = next(
        check
        for check in report.checks
        if check.code == "paper_total_matches"
    )
    text = _docx_text(docx)

    assert not any(
        issue.code == "paper_total_marks_mismatch"
        for issue in readiness.issues
    )
    assert marks_check.passed is True
    assert "الدرجة: 80" in text
    assert "عربية معتمدة" in text


def test_scaled_policy_records_raw_and_final_marks() -> None:
    project = _project(
        policy=MarksPolicy.scale_to_declared,
        translation_status=FullExamTranslationAcceptanceStatus.accepted,
    )

    readiness = build_project_readiness_report(project)
    docx = build_project_docx_bytes(project)
    text = _docx_text(docx)

    assert not any(
        issue.code == "paper_total_marks_mismatch"
        for issue in readiness.issues
    )
    assert "20 (محولة من 80)" in text
    assert "المجموع الخام" in text
    assert "80" in text


def test_unaccepted_translation_is_not_labeled_clean_arabic() -> None:
    project = _project(
        policy=MarksPolicy.use_question_total,
        translation_status=FullExamTranslationAcceptanceStatus.needs_review,
    )

    docx = build_project_docx_bytes(project)
    pdf = build_project_pdf_bytes(project)
    docx_text = _docx_text(docx)
    pdf_text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(pdf)).pages
    )

    assert "مسودة ترجمة تحتاج مراجعة" in docx_text
    assert "عربية نظيفة" not in docx_text
    assert "عربية نظيفة" not in pdf_text


def test_student_exports_hide_attachment_heading() -> None:
    project = _project(
        policy=MarksPolicy.use_question_total,
        translation_status=FullExamTranslationAcceptanceStatus.accepted,
    )

    docx = build_project_docx_bytes(project)
    pdf = build_project_pdf_bytes(project)
    docx_text = _docx_text(docx)
    pdf_text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(pdf)).pages
    )

    assert "مرفقات السؤال" not in docx_text
    assert "مرفقات السؤال" not in pdf_text


def test_docx_portrait_asset_is_bounded_without_distortion() -> None:
    project = _project(
        policy=MarksPolicy.use_question_total,
        translation_status=FullExamTranslationAcceptanceStatus.accepted,
    )

    document = Document(BytesIO(build_project_docx_bytes(project)))
    question_shapes = list(document.inline_shapes)

    assert len(question_shapes) == 1
    width_inches = question_shapes[0].width.inches
    height_inches = question_shapes[0].height.inches
    assert width_inches <= 3.55
    assert height_inches <= 2.65
    assert abs((width_inches / height_inches) - (400 / 1200)) < 0.02
