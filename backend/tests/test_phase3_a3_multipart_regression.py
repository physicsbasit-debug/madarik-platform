from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfReader

from app.main import app
from app.models.project import (
    OutputMode,
    ProjectMetadata,
    ProjectSession,
    QuestionItem,
)
from app.services.question_parts import parse_question_parts


client = TestClient(app)


HIERARCHICAL_SOURCE = (
    "(d) State whether the results agree. "
    "(e) "
    "(i) State the unit of force. [1] "
    "(ii) Calculate the resultant force. [2]"
)


def _docx_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    text_blocks = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_blocks.extend(
                    paragraph.text
                    for paragraph in cell.paragraphs
                )

    return "\n".join(text_blocks)


def _pdf_text(pdf_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(pdf_bytes))
    return "\n".join(
        page.extract_text() or ""
        for page in reader.pages
    )


def _parts_by_label(question: dict) -> dict[str, dict]:
    return {
        part["label"]: part
        for part in question["parts"]
    }


def test_phase3_a3_parser_and_legacy_payload_regression_lock() -> None:
    parts = parse_question_parts(HIERARCHICAL_SOURCE)

    assert [part.label for part in parts] == [
        "(d)",
        "(e)",
        "(i)",
        "(ii)",
    ]
    assert parts[1].original_text == ""
    assert parts[2].parent_id == parts[1].id
    assert parts[3].parent_id == parts[1].id
    assert [part.marks for part in parts] == [
        None,
        None,
        1,
        2,
    ]

    legacy_project = ProjectSession.model_validate(
        {
            "questions": [
                {
                    "id": "legacy-question",
                    "original_number": "1",
                    "original_text": "Legacy multipart question",
                    "translated_text": "",
                    "order_index": 1,
                    "parts": [
                        {
                            "id": "legacy-part",
                            "label": "(a)",
                            "original_text": "State the unit.",
                            "translated_text": "اذكر الوحدة.",
                            "marks": 1,
                            "order_index": 1,
                        }
                    ],
                }
            ]
        }
    )

    assert legacy_project.questions[0].parts[0].parent_id is None


def test_phase3_a3_full_hierarchical_question_workflow_regression() -> None:
    parsed_parts = parse_question_parts(HIERARCHICAL_SOURCE)
    accepted_parts = [
        part.model_copy(update={"marks": 3})
        if part.label == "(e)"
        else part
        for part in parsed_parts
    ]

    snapshot = ProjectSession(
        metadata=ProjectMetadata(
            paper_title="Phase 3-A3 regression paper",
            subject="الفيزياء",
            output_mode=OutputMode.bilingual,
        ),
        questions=[
            QuestionItem(
                id="phase3-a3-question",
                original_number="1",
                original_text=(
                    "RAW HIERARCHICAL QUESTION MUST NOT BE DUPLICATED"
                ),
                translated_text="",
                marks=1,
                order_index=1,
                parts=accepted_parts,
            )
        ],
    )

    imported_response = client.post(
        "/api/projects/import-snapshot",
        json=snapshot.model_dump(mode="json"),
    )
    assert imported_response.status_code == 201

    imported_project = imported_response.json()
    project_id = imported_project["id"]
    question = imported_project["questions"][0]
    question_id = question["id"]
    parts_by_label = _parts_by_label(question)
    parent_id = parts_by_label["(e)"]["id"]

    reordered_parts = []
    desired_order = {
        "(d)": 1,
        "(e)": 2,
        "(ii)": 3,
        "(i)": 4,
    }

    for part in question["parts"]:
        updated_part = dict(part)
        updated_part["order_index"] = desired_order[part["label"]]
        reordered_parts.append(updated_part)

    reordered_parts.sort(key=lambda part: part["order_index"])

    reorder_response = client.patch(
        f"/api/projects/{project_id}/questions/{question_id}",
        json={"parts": reordered_parts},
    )
    assert reorder_response.status_code == 200

    reordered_question = reorder_response.json()["questions"][0]
    assert [
        part["label"]
        for part in sorted(
            reordered_question["parts"],
            key=lambda part: part["order_index"],
        )
    ] == [
        "(d)",
        "(e)",
        "(ii)",
        "(i)",
    ]
    reordered_by_label = _parts_by_label(reordered_question)
    assert reordered_by_label["(ii)"]["parent_id"] == parent_id
    assert reordered_by_label["(i)"]["parent_id"] == parent_id

    saved_snapshot = client.get(
        f"/api/projects/{project_id}/snapshot"
    )
    assert saved_snapshot.status_code == 200

    round_trip_response = client.post(
        "/api/projects/import-snapshot",
        json=saved_snapshot.json(),
    )
    assert round_trip_response.status_code == 201

    round_trip_project = round_trip_response.json()
    round_trip_id = round_trip_project["id"]
    round_trip_question = round_trip_project["questions"][0]
    round_trip_by_label = _parts_by_label(round_trip_question)

    assert round_trip_by_label["(ii)"]["parent_id"] == (
        round_trip_by_label["(e)"]["id"]
    )
    assert round_trip_by_label["(i)"]["parent_id"] == (
        round_trip_by_label["(e)"]["id"]
    )
    assert round_trip_by_label["(ii)"]["order_index"] < (
        round_trip_by_label["(i)"]["order_index"]
    )

    translate_response = client.post(
        f"/api/projects/{round_trip_id}/translate-questions"
    )
    assert translate_response.status_code == 200

    translated_question = translate_response.json()["questions"][0]
    translated_by_label = _parts_by_label(translated_question)

    assert translated_by_label["(e)"]["translated_text"] == ""
    assert translated_by_label["(ii)"]["translated_text"].strip()
    assert translated_by_label["(i)"]["translated_text"].strip()
    assert translated_question["translated_text"].strip()
    assert translated_by_label["(ii)"]["parent_id"] == (
        translated_by_label["(e)"]["id"]
    )

    mismatch_readiness = client.get(
        f"/api/projects/{round_trip_id}/readiness"
    )
    assert mismatch_readiness.status_code == 200
    mismatch_body = mismatch_readiness.json()
    mismatch_codes = {
        issue["code"]
        for issue in mismatch_body["issues"]
    }

    assert mismatch_body["ready"] is True
    assert mismatch_body["total_marks"] == 1
    assert "question_parts_marks_mismatch" in mismatch_codes

    pre_adoption_docx = client.post(
        f"/api/projects/{round_trip_id}/export/docx"
    )
    assert pre_adoption_docx.status_code == 200
    assert pre_adoption_docx.content[:2] == b"PK"

    adopt_response = client.patch(
        (
            f"/api/projects/{round_trip_id}/questions/"
            f"{translated_question['id']}"
        ),
        json={"marks": 3},
    )
    assert adopt_response.status_code == 200

    accepted_readiness = client.get(
        f"/api/projects/{round_trip_id}/readiness"
    )
    assert accepted_readiness.status_code == 200
    accepted_body = accepted_readiness.json()
    accepted_codes = {
        issue["code"]
        for issue in accepted_body["issues"]
    }

    assert accepted_body["ready"] is True
    assert accepted_body["total_marks"] == 3
    assert "question_parts_marks_mismatch" not in accepted_codes

    final_snapshot = client.get(
        f"/api/projects/{round_trip_id}/snapshot"
    )
    assert final_snapshot.status_code == 200

    final_import = client.post(
        "/api/projects/import-snapshot",
        json=final_snapshot.json(),
    )
    assert final_import.status_code == 201
    final_project_id = final_import.json()["id"]

    final_project = client.get(
        f"/api/projects/{final_project_id}"
    )
    assert final_project.status_code == 200
    final_question = final_project.json()["questions"][0]
    final_by_label = _parts_by_label(final_question)

    assert final_question["marks"] == 3
    assert final_by_label["(ii)"]["translated_text"].strip()
    assert final_by_label["(ii)"]["parent_id"] == (
        final_by_label["(e)"]["id"]
    )
    assert final_by_label["(ii)"]["order_index"] < (
        final_by_label["(i)"]["order_index"]
    )

    docx_response = client.post(
        f"/api/projects/{final_project_id}/export/docx"
    )
    assert docx_response.status_code == 200
    assert docx_response.content[:2] == b"PK"

    docx_text = _docx_text(docx_response.content)
    assert "السؤال 1 [3]" in docx_text
    assert "مجموع الدرجات: 3" in docx_text
    assert "(e) [3]" in docx_text
    assert "(ii) [2]" in docx_text
    assert "(i) [1]" in docx_text
    assert docx_text.index("(e) [3]") < docx_text.index("(ii) [2]")
    assert docx_text.index("(ii) [2]") < docx_text.index("(i) [1]")
    assert "RAW HIERARCHICAL QUESTION MUST NOT BE DUPLICATED" not in docx_text

    pdf_response = client.post(
        f"/api/projects/{final_project_id}/export/pdf"
    )
    assert pdf_response.status_code == 200
    assert pdf_response.content[:4] == b"%PDF"

    pdf_text = _pdf_text(pdf_response.content)
    assert "e) [3]" in pdf_text
    assert "ii) [2]" in pdf_text
    assert "i) [1]" in pdf_text
    assert pdf_text.index("e) [3]") < pdf_text.index("ii) [2]")
    assert pdf_text.index("ii) [2]") < pdf_text.index("i) [1]")
    assert "RAW HIERARCHICAL QUESTION MUST NOT BE DUPLICATED" not in pdf_text
