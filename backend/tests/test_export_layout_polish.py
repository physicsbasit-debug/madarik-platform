from __future__ import annotations

import base64
from io import BytesIO

from docx import Document
from docx.shared import Inches
from PIL import Image as PILImage

from app.models.project import (
    OutputMode,
    ProjectMetadata,
    ProjectSession,
    QuestionAssetInfo,
    QuestionItem,
)
from app.services.export import (
    _clean_export_text,
    _question_heading_text,
    build_project_docx_bytes,
)


def _test_png_base64() -> str:
    image = PILImage.new(
        "RGB",
        (120, 90),
        (255, 255, 255),
    )
    output = BytesIO()
    image.save(
        output,
        format="PNG",
    )

    return base64.b64encode(
        output.getvalue()
    ).decode("ascii")


def _build_project() -> ProjectSession:
    image_data = _test_png_base64()

    return ProjectSession(
        metadata=ProjectMetadata(
            paper_title="تدريب مترجم",
            school_name="مدرسة اختبار",
            subject="الفيزياء",
            grade="العاشر",
            semester="الفصل الأول",
            duration="45 دقيقة",
            teacher_name="معلم الاختبار",
            date="13-07-2026",
            total_marks="20",
            output_mode=OutputMode.bilingual,
        ),
        questions=[
            QuestionItem(
                id="question-1",
                original_number="1",
                original_text=(
                    "(b) Fig. 2.2 shows a man using "
                    "a golf club to hit a ball. "
                    "golf clubballFig. 2.2 "
                    "The ball has a mass of 0.046 kg."
                ),
                translated_text="ترجمة تجريبية",
                marks=2,
                order_index=1,
                attachments=[
                    QuestionAssetInfo(
                        id="asset-1",
                        name="figure.png",
                        size=len(
                            base64.b64decode(
                                image_data
                            )
                        ),
                        type="image/png",
                        data_base64=image_data,
                    )
                ],
            )
        ],
    )


def test_clean_export_text_collapses_duplicate_figure_caption() -> None:
    cleaned = _clean_export_text(
        "(b) Fig. 2.2 shows a man using "
        "a golf club to hit a ball. "
        "golf clubballFig. 2.2 "
        "The ball has a mass."
    )

    assert cleaned.count("Fig. 2.2") == 1
    assert "golf clubball" not in cleaned
    assert "The ball has a mass." in cleaned


def test_question_heading_uses_arabic_order() -> None:
    question = _build_project().questions[0]

    assert (
        _question_heading_text(
            1,
            question,
        )
        == "السؤال 1 [2]"
    )


def test_docx_uses_summary_table_and_smaller_image() -> None:
    payload = build_project_docx_bytes(
        _build_project()
    )

    document = Document(
        BytesIO(payload)
    )

    assert len(document.tables) == 2

    summary_cells = [
        cell.text
        for cell in document.tables[1].rows[0].cells
    ]

    assert any(
        "نوع النسخة" in text
        and "ثنائية اللغة" in text
        for text in summary_cells
    )
    assert any(
        "عدد الأسئلة" in text
        and "1" in text
        for text in summary_cells
    )
    assert any(
        "مجموع الدرجات" in text
        and "2" in text
        for text in summary_cells
    )

    paragraph_text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )

    assert "السؤال 1 [2]" in paragraph_text
    assert "1. السؤال" not in paragraph_text
    assert paragraph_text.count("Fig. 2.2") == 1
    assert "golf clubball" not in paragraph_text

    assert len(document.inline_shapes) == 1

    shape = document.inline_shapes[0]

    assert shape.width <= Inches(3.55)
    assert shape.height <= Inches(2.65)
