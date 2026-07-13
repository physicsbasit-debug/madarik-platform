from __future__ import annotations

import base64
from io import BytesIO

from docx import Document
from PIL import Image as PILImage

from app.models.project import (
    OutputMode,
    ProjectMetadata,
    ProjectSession,
    QuestionAssetInfo,
    QuestionItem,
)
from app.services.export import (
    _attachment_note_for_export,
    _clean_export_text,
    _split_export_blocks,
    build_project_docx_bytes,
)


def _test_png_base64() -> str:
    image = PILImage.new(
        "RGB",
        (12, 8),
        (255, 255, 255),
    )
    output = BytesIO()
    image.save(output, format="PNG")

    return base64.b64encode(
        output.getvalue()
    ).decode("ascii")


def _docx_text(payload: bytes) -> str:
    document = Document(BytesIO(payload))
    parts = [
        paragraph.text
        for paragraph in document.paragraphs
    ]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(
                    paragraph.text
                    for paragraph in cell.paragraphs
                )

    return "\n".join(parts)


def test_clean_export_text_removes_xml_controls_and_squares() -> None:
    cleaned = _clean_export_text(
        "The■ball\x00has\x0bmass."
    )

    assert cleaned == "The ball has mass."
    assert "\x00" not in cleaned
    assert "\x0b" not in cleaned
    assert "■" not in cleaned


def test_split_export_blocks_separates_question_parts() -> None:
    blocks = _split_export_blocks(
        "(a) First part. (b) Second part. "
        "(i) Calculation. (ii) Explanation."
    )

    assert blocks == [
        "(a) First part.",
        "(b) Second part.",
        "(i) Calculation.",
        "(ii) Explanation.",
    ]


def test_stale_attachment_note_is_hidden_when_asset_exists() -> None:
    question = QuestionItem(
        id="question-1",
        original_number="1",
        original_text="Original",
        translated_text="ترجمة",
        order_index=1,
        attachment_note=(
            "لم يتم ربط الصور والجداول بعد. "
            "هذه الوظيفة مؤجلة."
        ),
        attachments=[
            QuestionAssetInfo(
                id="asset-1",
                name="figure.png",
                size=10,
                type="image/png",
                data_base64=_test_png_base64(),
            )
        ],
    )

    assert _attachment_note_for_export(question) == ""


def test_docx_export_accepts_ocr_control_characters() -> None:
    image_data = _test_png_base64()

    project = ProjectSession(
        metadata=ProjectMetadata(
            paper_title="ورقة\x00 تدريبية",
            output_mode=OutputMode.bilingual,
        ),
        questions=[
            QuestionItem(
                id="question-1",
                original_number="1",
                original_text=(
                    "(a) The■ball\x0bhas mass. "
                    "(b) Calculate <force> & time. "
                    "(i) Show working."
                ),
                translated_text=(
                    "ترجمة تحتاج مراجعة\x00"
                ),
                order_index=1,
                attachment_note=(
                    "لم يتم ربط الصور والجداول بعد. "
                    "هذه الوظيفة مؤجلة."
                ),
                attachments=[
                    QuestionAssetInfo(
                        id="asset-1",
                        name="figure.png",
                        size=len(
                            base64.b64decode(image_data)
                        ),
                        type="image/png",
                        data_base64=image_data,
                    )
                ],
            )
        ],
    )

    payload = build_project_docx_bytes(project)
    exported_text = _docx_text(payload)

    assert payload.startswith(b"PK")
    assert "The ball has mass." in exported_text
    assert "(b) Calculate <force> & time." in exported_text
    assert "■" not in exported_text
    assert "\x00" not in exported_text
    assert "\x0b" not in exported_text
    assert "لم يتم ربط الصور والجداول" not in exported_text
