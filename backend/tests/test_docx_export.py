from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from fastapi.testclient import TestClient

from app.main import app
from app.models.project import (
    OutputMode,
    ProjectMetadata,
    ProjectSession,
    QuestionItem,
    QuestionPart,
    QuestionStatus,
)
from app.services.export import build_project_docx_bytes

client = TestClient(app)


def _read_docx_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def test_build_project_docx_arabic_export_filters_deleted_questions():
    project = ProjectSession(
        metadata=ProjectMetadata(
            school_name="مدرسة الباسط للتعليم الأساسي",
            subject="الفيزياء",
            grade="العاشر",
            paper_title="تدريب القوى",
            teacher_name="أ. وليد الهنائي",
            output_mode=OutputMode.arabic,
        ),
        questions=[
            QuestionItem(
                id="q-1",
                original_number="1",
                original_text="State the unit of force. [1]",
                translated_text="اذكر وحدة القوة. [1]",
                marks=1,
                detected_marks=1,
                order_index=2,
            ),
            QuestionItem(
                id="q-2",
                original_number="2",
                original_text="Explain why the object accelerates. [2]",
                translated_text="فسّر لماذا يتسارع الجسم. [2]",
                marks=2,
                detected_marks=2,
                order_index=1,
            ),
            QuestionItem(
                id="q-3",
                original_number="3",
                original_text="Deleted question",
                translated_text="سؤال محذوف",
                marks=1,
                detected_marks=1,
                order_index=3,
                status=QuestionStatus.deleted,
            ),
        ],
    )

    docx_bytes = build_project_docx_bytes(project)

    assert docx_bytes.startswith(b"PK")
    text = _read_docx_text(docx_bytes)
    assert "تدريب القوى" in text
    assert "مدرسة الباسط" in text
    assert "فسّر لماذا يتسارع الجسم" in text
    assert "اذكر وحدة القوة" in text
    assert "سؤال محذوف" not in text
    assert "السؤال 1 [2]" in text
    assert "السؤال 2 [1]" in text


def test_build_project_docx_bilingual_export_includes_original_and_translation():
    project = ProjectSession(
        metadata=ProjectMetadata(
            paper_title="Cells practice",
            subject="الأحياء",
            output_mode=OutputMode.bilingual,
        ),
        questions=[
            QuestionItem(
                id="q-1",
                original_number="1",
                original_text="State the function of the cell membrane. [1]",
                translated_text="اذكر وظيفة غشاء الخلية. [1]",
                marks=1,
                detected_marks=1,
                order_index=1,
            )
        ],
    )

    text = _read_docx_text(build_project_docx_bytes(project))

    assert "State the function of the cell membrane" in text
    assert "اذكر وظيفة غشاء الخلية" in text
    assert "ثنائية اللغة" in text


def test_build_project_docx_exports_question_parts_as_structured_blocks():
    project = ProjectSession(
        metadata=ProjectMetadata(
            paper_title="Multipart practice",
            subject="الفيزياء",
            output_mode=OutputMode.bilingual,
        ),
        questions=[
            QuestionItem(
                id="q-parts",
                original_number="4",
                original_text="RAW MULTIPART TEXT MUST NOT BE DUPLICATED",
                translated_text="ترجمة مجمعة لا ينبغي تكرارها",
                marks=None,
                detected_marks=None,
                order_index=1,
                parts=[
                    QuestionPart(
                        id="part-b",
                        label="(b)",
                        original_text="Calculate the resultant force.",
                        translated_text="احسب القوة المحصلة.",
                        marks=2,
                        order_index=2,
                    ),
                    QuestionPart(
                        id="part-a",
                        label="(a)",
                        original_text="State the unit of force.",
                        translated_text="اذكر وحدة القوة.",
                        marks=1,
                        order_index=1,
                    ),
                ],
            )
        ],
    )

    text = _read_docx_text(build_project_docx_bytes(project))

    assert "السؤال 1 [3]" in text
    assert "(a) [1]" in text
    assert "(b) [2]" in text
    assert text.index("(a) [1]") < text.index("(b) [2]")
    assert "State the unit of force." in text
    assert "اذكر وحدة القوة." in text
    assert "Calculate the resultant force." in text
    assert "احسب القوة المحصلة." in text
    assert "RAW MULTIPART TEXT MUST NOT BE DUPLICATED" not in text
    assert "ترجمة مجمعة لا ينبغي تكرارها" not in text
    assert "مجموع الدرجات: 3" in text



def test_docx_export_indents_child_parts_and_avoids_double_counting_marks():
    project = ProjectSession(
        metadata=ProjectMetadata(
            paper_title="Hierarchical multipart practice",
            subject="الفيزياء",
            output_mode=OutputMode.bilingual,
        ),
        questions=[
            QuestionItem(
                id="q-hierarchy-docx",
                original_number="5",
                original_text="RAW HIERARCHICAL TEXT",
                translated_text="",
                order_index=1,
                parts=[
                    QuestionPart(
                        id="part-e",
                        label="(e)",
                        original_text="",
                        translated_text="",
                        marks=3,
                        order_index=1,
                    ),
                    QuestionPart(
                        id="part-i",
                        label="(i)",
                        original_text="State the unit.",
                        translated_text="اذكر الوحدة.",
                        marks=1,
                        parent_id="part-e",
                        order_index=2,
                    ),
                    QuestionPart(
                        id="part-ii",
                        label="(ii)",
                        original_text="Calculate the value.",
                        translated_text="احسب القيمة.",
                        marks=2,
                        parent_id="part-e",
                        order_index=3,
                    ),
                ],
            )
        ],
    )

    document = Document(BytesIO(build_project_docx_bytes(project)))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    parent_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "(e) [3]"
    )
    child_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "(i) [1]"
    )

    assert "السؤال 1 [3]" in text
    assert "[Original text unavailable]" not in text
    assert child_heading.paragraph_format.right_indent > (
        parent_heading.paragraph_format.right_indent
    )


def test_docx_mixed_direction_headings_keep_stable_alignment():
    project = ProjectSession(
        metadata=ProjectMetadata(
            paper_title="تدريب مترجم من ورقة اختبار دولية",
            subject="الفيزياء",
            output_mode=OutputMode.bilingual,
        ),
        questions=[
            QuestionItem(
                id="q-parts-direction",
                original_number="4",
                original_text="RAW MULTIPART TEXT",
                translated_text="ترجمة مجمعة",
                order_index=1,
                parts=[
                    QuestionPart(
                        id="part-i",
                        label="(i)",
                        original_text="State the unit of force.",
                        translated_text="اذكر وحدة القوة.",
                        marks=1,
                        order_index=1,
                    )
                ],
            )
        ],
    )

    document = Document(BytesIO(build_project_docx_bytes(project)))

    title = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "تدريب مترجم من ورقة اختبار دولية"
    )
    question_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "السؤال 1 [1]"
    )
    part_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "(i) [1]"
    )

    assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert question_heading.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert part_heading.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert question_heading._p.pPr.find(qn("w:bidi")) is None
    assert part_heading._p.pPr.find(qn("w:bidi")) is None

def test_export_docx_endpoint_returns_downloadable_word_file():
    create_response = client.post("/api/projects")
    project_id = create_response.json()["id"]

    demo_response = client.post(f"/api/projects/{project_id}/demo-content")
    assert demo_response.status_code == 200

    translate_response = client.post(f"/api/projects/{project_id}/translate-questions")
    assert translate_response.status_code == 200

    response = client.post(f"/api/projects/{project_id}/export/docx")

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "attachment" in response.headers["content-disposition"]
    assert response.content.startswith(b"PK")
    assert "منصة مدارك" in _read_docx_text(response.content)


def test_export_docx_endpoint_rejects_empty_active_questions():
    create_response = client.post("/api/projects")
    project_id = create_response.json()["id"]

    response = client.post(f"/api/projects/{project_id}/export/docx")

    assert response.status_code == 400
    assert "لا توجد أسئلة" in response.json()["detail"]
