from __future__ import annotations

from io import BytesIO
import base64
import os
import re
from typing import Iterable
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

import arabic_reshaper
from bidi.algorithm import get_display

from app.models.project import (
    OutputMode,
    ProjectSession,
    QuestionItem,
    QuestionPart,
    QuestionStatus,
)

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME_TYPE = "application/pdf"

ARABIC_FONT_NAME = "MadarikArabic"
ENGLISH_FONT_NAME = "Helvetica"

_STALE_ATTACHMENT_NOTE_SNIPPETS = (
    "لم يتم ربط الصور والجداول",
    "هذه الوظيفة مؤجلة",
)

_SUBQUESTION_MARKER_RE = re.compile(
    r"(?<!^)\s+"
    r"(?=(?:"
    r"\([a-z]\)"
    r"|\([ivxlcdm]+\)"
    r"|\(\d+\)"
    r")\s*)",
    re.IGNORECASE,
)


_DUPLICATE_FIGURE_REFERENCE_RE = re.compile(
    r"("
    r"Fig\.\s*(?P<number>\d+(?:\.\d+)*)\b"
    r".{0,260}?[.!?]"
    r")"
    r"\s*"
    r"[A-Za-z][A-Za-z \-]{1,70}?"
    r"Fig\.\s*(?P=number)\b",
    re.IGNORECASE | re.DOTALL,
)


def _collapse_duplicate_figure_references(
    value: str,
) -> str:
    """Remove duplicated figure captions and OCR image-label debris."""

    current = value

    for _ in range(3):
        updated = _DUPLICATE_FIGURE_REFERENCE_RE.sub(
            r"\1 ",
            current,
        )

        if updated == current:
            break

        current = updated

    return current


def _is_xml_compatible_character(character: str) -> bool:
    codepoint = ord(character)

    return (
        codepoint in (0x09, 0x0A, 0x0D)
        or 0x20 <= codepoint <= 0xD7FF
        or 0xE000 <= codepoint <= 0xFFFD
        or 0x10000 <= codepoint <= 0x10FFFF
    )



def _clean_export_text(
    value: object | None,
) -> str:
    """Clean OCR text for DOCX XML and ReportLab paragraphs."""

    raw = str(value or "")
    raw = raw.replace("\r\n", "\n").replace("\r", "\n")

    cleaned = "".join(
        character
        if _is_xml_compatible_character(character)
        else " "
        for character in raw
    )

    for separator in (
        "\u25a0",
        "\u25aa",
        "\u25a1",
        "\ufffd",
    ):
        cleaned = cleaned.replace(separator, " ")

    cleaned = cleaned.replace("\u00a0", " ")

    # OCR sometimes joins two sentences as: quantity.mass
    cleaned = re.sub(
        r"(?<=[A-Za-z0-9])\.(?=[A-Za-z])",
        ". ",
        cleaned,
    )

    # OCR sometimes glues image labels to a repeated figure caption:
    # "... ball. golf clubballFig. 2.2 The ball ..."
    cleaned = _collapse_duplicate_figure_references(
        cleaned,
    )

    normalized_lines: list[str] = []
    previous_was_blank = False

    for raw_line in cleaned.split("\n"):
        line = re.sub(
            r"[ \t]+",
            " ",
            raw_line,
        ).strip()

        if line:
            normalized_lines.append(line)
            previous_was_blank = False
        elif normalized_lines and not previous_was_blank:
            normalized_lines.append("")
            previous_was_blank = True

    return "\n".join(normalized_lines).strip()

def _split_export_blocks(value: object | None) -> list[str]:
    """Split OCR text into readable question-part paragraphs."""

    cleaned = _clean_export_text(value)

    if not cleaned:
        return []

    separated = _SUBQUESTION_MARKER_RE.sub(
        "\n",
        cleaned,
    )

    return [
        block.strip()
        for block in separated.splitlines()
        if block.strip()
    ]


def _attachment_note_for_export(
    question: QuestionItem,
) -> str:
    note = _clean_export_text(question.attachment_note)

    if not note:
        return ""

    if question.attachments and any(
        snippet in note
        for snippet in _STALE_ATTACHMENT_NOTE_SNIPPETS
    ):
        return ""

    return note


def _add_docx_text_blocks(
    document: Document,
    text: object | None,
    *,
    rtl: bool,
    fallback: str,
    font_size: float,
    indent_level: int = 0,
) -> None:
    blocks = _split_export_blocks(text) or [fallback]

    for block in blocks:
        paragraph = document.add_paragraph()
        _set_paragraph_bidi(paragraph, rtl=rtl)

        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15

        safe_indent_level = max(indent_level, 0)

        if rtl:
            if safe_indent_level:
                paragraph.paragraph_format.right_indent = Inches(
                    0.2 * safe_indent_level
                )
        else:
            paragraph.paragraph_format.left_indent = Inches(
                0.2 + (0.2 * safe_indent_level)
            )

        run = paragraph.add_run(
            _clean_export_text(block),
        )
        run.font.size = Pt(font_size)
        _set_run_rtl(run, rtl=rtl)



def _set_paragraph_bidi(paragraph, *, rtl: bool = True) -> None:
    """Set paragraph direction for Arabic/English blocks."""

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        p_pr.append(OxmlElement("w:bidi"))
    if not rtl and bidi is not None:
        p_pr.remove(bidi)


def _set_run_rtl(run, *, rtl: bool = True) -> None:
    r_pr = run._r.get_or_add_rPr()
    rtl_el = r_pr.find(qn("w:rtl"))
    if rtl and rtl_el is None:
        r_pr.append(OxmlElement("w:rtl"))
    if not rtl and rtl_el is not None:
        r_pr.remove(rtl_el)



def _set_cell_text(
    cell,
    label: str,
    value: str,
) -> None:
    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    paragraph = cell.paragraphs[0]
    paragraph.clear()
    _set_paragraph_bidi(paragraph, rtl=True)

    safe_label = _clean_export_text(label)
    safe_value = (
        _clean_export_text(value)
        or "__________"
    )

    label_run = paragraph.add_run(
        f"{safe_label}: "
    )
    label_run.bold = True
    _set_run_rtl(label_run, rtl=True)

    value_run = paragraph.add_run(safe_value)
    _set_run_rtl(value_run, rtl=True)

def _set_table_bidi(table) -> None:
    tbl_pr = table._tbl.tblPr
    bidi = tbl_pr.find(qn("w:bidiVisual"))
    if bidi is None:
        tbl_pr.append(OxmlElement("w:bidiVisual"))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.start_type = WD_SECTION_START.NEW_PAGE

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(12)


def _logo_bytes(project: ProjectSession) -> bytes | None:
    if not project.school_logo or not project.school_logo.data_base64:
        return None
    try:
        return base64.b64decode(project.school_logo.data_base64)
    except Exception:
        return None


def _asset_bytes(data_base64: str) -> bytes | None:
    if not data_base64:
        return None
    try:
        return base64.b64decode(data_base64)
    except Exception:
        return None


def _add_docx_logo(document: Document, project: ProjectSession) -> None:
    logo_bytes = _logo_bytes(project)
    if not logo_bytes:
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = paragraph.add_run()
        run.add_picture(BytesIO(logo_bytes), width=Inches(0.9))
    except Exception:
        # A broken or unsupported logo must not block exporting the teacher's paper.
        paragraph.clear()



def _add_title(
    document: Document,
    project: ProjectSession,
) -> None:
    title = document.add_paragraph()
    _set_paragraph_bidi(title, rtl=True)
    # Set centering after bidi because _set_paragraph_bidi assigns a
    # direction-based alignment by default.
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run(
        _clean_export_text(
            project.metadata.paper_title
            or "ورقة تدريبية مترجمة"
        )
    )
    title_run.bold = True
    title_run.font.size = Pt(17)
    _set_run_rtl(title_run, rtl=True)

    subtitle = document.add_paragraph()
    _set_paragraph_bidi(subtitle, rtl=True)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run("منصة مدارك")
    subtitle_run.font.size = Pt(10)
    _set_run_rtl(subtitle_run, rtl=True)


def _add_metadata_table(
    document: Document,
    project: ProjectSession,
    question_count: int,
    marks_total: int,
) -> None:
    metadata = project.metadata

    table = document.add_table(
        rows=4,
        cols=2,
    )
    table.style = "Table Grid"
    _set_table_bidi(table)

    rows = [
        (
            ("اسم المدرسة", metadata.school_name),
            ("المادة", metadata.subject),
        ),
        (
            ("الصف", metadata.grade),
            ("الفصل الدراسي", metadata.semester),
        ),
        (
            ("الزمن", metadata.duration),
            (
                "الدرجة",
                metadata.total_marks
                or str(marks_total),
            ),
        ),
        (
            ("اسم المعلم", metadata.teacher_name),
            ("التاريخ", metadata.date),
        ),
    ]

    for row, row_data in zip(
        table.rows,
        rows,
    ):
        _set_cell_text(
            row.cells[0],
            row_data[0][0],
            row_data[0][1],
        )
        _set_cell_text(
            row.cells[1],
            row_data[1][0],
            row_data[1][1],
        )

    mode = (
        "ثنائية اللغة"
        if metadata.output_mode == OutputMode.bilingual
        else "عربية نظيفة"
    )

    summary_table = document.add_table(
        rows=1,
        cols=3,
    )
    summary_table.style = "Table Grid"
    _set_table_bidi(summary_table)

    summary_items = [
        ("نوع النسخة", mode),
        ("عدد الأسئلة", str(question_count)),
        ("مجموع الدرجات", str(marks_total)),
    ]

    for cell, (label, value) in zip(
        summary_table.rows[0].cells,
        summary_items,
    ):
        _set_cell_text(
            cell,
            label,
            value,
        )

def _active_questions(project: ProjectSession) -> list[QuestionItem]:
    return sorted(
        [question for question in project.questions if question.status != QuestionStatus.deleted],
        key=lambda question: question.order_index,
    )


def _question_marks_label(question: QuestionItem) -> str:
    marks = _question_total_marks(question)
    if marks is None:
        return ""
    return f" [{marks}]"


def _question_part_depth(
    part: QuestionPart,
    parts: list[QuestionPart],
) -> int:
    """Return a cycle-safe hierarchy depth for one structured part."""

    parts_by_id = {
        item.id: item
        for item in parts
    }
    depth = 0
    current = part
    visited = {part.id}

    while current.parent_id:
        parent = parts_by_id.get(current.parent_id)

        if parent is None or parent.id in visited:
            break

        visited.add(parent.id)
        depth += 1
        current = parent

    return depth


def _question_total_marks(question: QuestionItem) -> int | None:
    """Return explicit marks or a hierarchy-safe total.

    Parent marks often summarize their children. Summing every row would then
    double-count the question, so each root branch prefers descendant marks and
    falls back to the root mark only when the branch has no marked children.
    """

    if question.marks is not None:
        return question.marks

    if not question.parts:
        return None

    parts_by_id = {
        part.id: part
        for part in question.parts
    }
    children: dict[str, list[QuestionPart]] = {}

    for part in question.parts:
        if part.parent_id in parts_by_id:
            children.setdefault(part.parent_id, []).append(part)

    roots = [
        part
        for part in question.parts
        if part.parent_id not in parts_by_id
    ]

    def branch_marks(
        part: QuestionPart,
        visited: set[str],
    ) -> int | None:
        if part.id in visited:
            return part.marks

        next_visited = visited | {part.id}
        child_values = [
            branch_marks(child, next_visited)
            for child in children.get(part.id, [])
        ]
        concrete_child_values = [
            value
            for value in child_values
            if value is not None
        ]

        if concrete_child_values:
            return sum(concrete_child_values)

        return part.marks

    totals = [
        branch_marks(root, set())
        for root in roots
    ]
    concrete_totals = [
        value
        for value in totals
        if value is not None
    ]

    return sum(concrete_totals) if concrete_totals else None


def _part_heading_text(part: QuestionPart) -> str:
    label = _clean_export_text(part.label) or "(جزء)"
    marks = f" [{part.marks}]" if part.marks is not None else ""
    return f"{label}{marks}"


def _sorted_question_parts(question: QuestionItem) -> list[QuestionPart]:
    return sorted(question.parts, key=lambda part: part.order_index)


def _question_heading_text(
    number: int,
    question: QuestionItem,
) -> str:
    return (
        f"السؤال {number}"
        f"{_question_marks_label(question)}"
    )


def _add_question_number(
    document: Document,
    number: int,
    question: QuestionItem,
) -> None:
    paragraph = document.add_paragraph()
    # The heading mixes Arabic with Latin digits/brackets. Keeping the
    # paragraph LTR but right-aligned prevents LibreOffice/Word from
    # moving the whole heading to the left margin or reversing [marks].
    _set_paragraph_bidi(
        paragraph,
        rtl=False,
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)

    run = paragraph.add_run(
        _question_heading_text(
            number,
            question,
        )
    )
    run.bold = True
    run.font.size = Pt(13)
    _set_run_rtl(
        run,
        rtl=True,
    )

def _add_arabic_text(
    document: Document,
    text: str,
    *,
    indent_level: int = 0,
) -> None:
    _add_docx_text_blocks(
        document,
        text,
        rtl=True,
        fallback="[ترجمة تحتاج مراجعة]",
        font_size=12,
        indent_level=indent_level,
    )


def _add_english_text(
    document: Document,
    text: str,
    *,
    indent_level: int = 0,
) -> None:
    _add_docx_text_blocks(
        document,
        text,
        rtl=False,
        fallback="[Original text unavailable]",
        font_size=10.5,
        indent_level=indent_level,
    )


def _add_docx_part_heading(
    document: Document,
    part: QuestionPart,
    *,
    depth: int = 0,
) -> None:
    paragraph = document.add_paragraph()
    # Part labels such as (i) [1] are LTR tokens. Use an LTR paragraph
    # aligned to the right so their parentheses and marks stay intact.
    _set_paragraph_bidi(paragraph, rtl=False)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.right_indent = Inches(
        0.18 + (0.28 * max(depth, 0))
    )
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)

    run = paragraph.add_run(_part_heading_text(part))
    run.bold = True
    run.font.size = Pt(11.5)
    _set_run_rtl(run, rtl=False)


def _add_docx_question_parts(
    document: Document,
    project: ProjectSession,
    question: QuestionItem,
) -> None:
    for part in _sorted_question_parts(question):
        depth = _question_part_depth(
            part,
            question.parts,
        )
        _add_docx_part_heading(
            document,
            part,
            depth=depth,
        )

        has_part_text = bool(
            part.original_text.strip()
            or part.translated_text.strip()
        )

        if not has_part_text:
            continue

        if project.metadata.output_mode == OutputMode.bilingual:
            _add_english_text(
                document,
                part.original_text,
                indent_level=depth,
            )
            _add_arabic_text(
                document,
                part.translated_text or "[تحتاج ترجمة]",
                indent_level=depth,
            )
        else:
            _add_arabic_text(
                document,
                part.translated_text or part.original_text,
                indent_level=depth,
            )



def _add_docx_question_assets(
    document: Document,
    question: QuestionItem,
) -> None:
    if not question.attachments:
        return

    label = document.add_paragraph()
    _set_paragraph_bidi(
        label,
        rtl=True,
    )

    label.paragraph_format.space_before = Pt(5)
    label.paragraph_format.space_after = Pt(3)

    label_run = label.add_run(
        "مرفقات السؤال:"
    )
    label_run.bold = True
    label_run.font.size = Pt(10)
    _set_run_rtl(
        label_run,
        rtl=True,
    )

    for asset in question.attachments:
        asset_bytes = _asset_bytes(
            asset.data_base64,
        )

        if not asset_bytes:
            continue

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)

        try:
            run = paragraph.add_run()
            run.add_picture(
                BytesIO(asset_bytes),
                width=Inches(3.8),
            )
        except Exception:
            paragraph.clear()
            _set_paragraph_bidi(
                paragraph,
                rtl=True,
            )

            broken = paragraph.add_run(
                "تعذر إدراج المرفق: "
                + _clean_export_text(asset.name)
            )
            broken.font.size = Pt(9)
            broken.italic = True
            _set_run_rtl(
                broken,
                rtl=True,
            )

def _add_questions(
    document: Document,
    project: ProjectSession,
    questions: list[QuestionItem],
) -> None:
    for display_number, question in enumerate(
        questions,
        start=1,
    ):
        _add_question_number(
            document,
            display_number,
            question,
        )

        if question.parts:
            _add_docx_question_parts(
                document,
                project,
                question,
            )
        elif (
            project.metadata.output_mode
            == OutputMode.bilingual
        ):
            _add_english_text(
                document,
                question.original_text,
            )
            _add_arabic_text(
                document,
                question.translated_text
                or "[تحتاج ترجمة]",
            )
        else:
            _add_arabic_text(
                document,
                question.translated_text
                or question.original_text,
            )

        _add_docx_question_assets(
            document,
            question,
        )

        attachment_note = (
            _attachment_note_for_export(question)
        )

        if attachment_note:
            note = document.add_paragraph()
            _set_paragraph_bidi(note, rtl=True)

            note_run = note.add_run(
                f"ملاحظة مرفق: {attachment_note}"
            )
            note_run.italic = True
            note_run.font.size = Pt(10)
            _set_run_rtl(note_run, rtl=True)

def _add_footer_note(document: Document) -> None:
    paragraph = document.add_paragraph()
    _set_paragraph_bidi(paragraph, rtl=True)
    run = paragraph.add_run("أُنشئت هذه الورقة عبر منصة مدارك. يُوصى بمراجعة المعلم للترجمة قبل الاستخدام الصفي.")
    run.font.size = Pt(9)
    run.italic = True
    _set_run_rtl(run, rtl=True)


def build_project_docx_bytes(project: ProjectSession) -> bytes:
    """Build a Phase 1-F1 RTL DOCX for the active project questions."""

    questions = _active_questions(project)
    if not questions:
        raise ValueError("لا توجد أسئلة نشطة قابلة للتصدير.")

    document = Document()
    _configure_document(document)

    marks_total = sum(_question_total_marks(question) or 0 for question in questions)
    _add_docx_logo(document, project)
    _add_title(document, project)
    _add_metadata_table(document, project, question_count=len(questions), marks_total=marks_total)
    _add_questions(document, project, questions)
    _add_footer_note(document)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def safe_docx_filename(project: ProjectSession) -> str:
    title = project.metadata.paper_title or "madarik_export"
    subject = project.metadata.subject or "paper"
    raw = f"madarik_{subject}_{title}"
    normalized = re.sub(r"[^A-Za-z0-9_\-]+", "_", raw).strip("_")
    if not normalized:
        normalized = "madarik_export"
    return f"{normalized[:90]}.docx"


# -----------------------------
# Phase 1-F2: PDF export
# -----------------------------


def _font_candidates() -> Iterable[str]:
    yield "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    yield "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"
    yield "/Library/Fonts/Arial Unicode.ttf"
    yield "C:/Windows/Fonts/arial.ttf"


def _ensure_pdf_font() -> str:
    """Register a Unicode-capable font when available.

    GitHub Actions Ubuntu images normally include DejaVu Sans. If no Unicode font
    is available, the PDF still builds with Helvetica, but Arabic glyph coverage
    will be limited. This fallback keeps the endpoint testable without bundling
    font files in the repository.
    """

    if ARABIC_FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return ARABIC_FONT_NAME

    for path in _font_candidates():
        if os.path.exists(path):
            pdfmetrics.registerFont(TTFont(ARABIC_FONT_NAME, path))
            return ARABIC_FONT_NAME

    return ENGLISH_FONT_NAME


def _shape_arabic(text: str) -> str:
    """Shape Arabic text for ReportLab's left-to-right drawing model."""

    if not text:
        return ""
    reshaped = arabic_reshaper.reshape(text)
    return get_display(reshaped)



def _pdf_paragraph(
    text: str,
    style: ParagraphStyle,
    *,
    rtl: bool = True,
) -> Paragraph:
    cleaned = _clean_export_text(text)

    prepared = (
        _shape_arabic(cleaned)
        if rtl
        else cleaned
    )

    safe = xml_escape(prepared).replace(
        "\n",
        "<br/>",
    )

    return Paragraph(
        safe or " ",
        style,
    )


def _add_pdf_text_blocks(
    story: list,
    text: object | None,
    style: ParagraphStyle,
    *,
    rtl: bool,
    fallback: str,
    indent_level: int = 0,
) -> None:
    blocks = _split_export_blocks(text) or [fallback]
    safe_indent_level = max(indent_level, 0)

    if safe_indent_level:
        indentation = safe_indent_level * 18
        style = ParagraphStyle(
            f"{style.name}Depth{safe_indent_level}",
            parent=style,
            rightIndent=(
                style.rightIndent + indentation
                if rtl
                else style.rightIndent
            ),
            leftIndent=(
                style.leftIndent + indentation
                if not rtl
                else style.leftIndent
            ),
        )

    for block in blocks:
        story.append(
            _pdf_paragraph(
                block,
                style,
                rtl=rtl,
            )
        )

def _pdf_styles() -> dict[str, ParagraphStyle]:
    font = _ensure_pdf_font()
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "MadarikTitle",
            parent=base["Title"],
            fontName=font,
            fontSize=18,
            leading=24,
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "subtitle": ParagraphStyle(
            "MadarikSubtitle",
            parent=base["Normal"],
            fontName=font,
            fontSize=10,
            leading=14,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#475569"),
            spaceAfter=12,
        ),
        "body_ar": ParagraphStyle(
            "MadarikBodyArabic",
            parent=base["Normal"],
            fontName=font,
            fontSize=11,
            leading=17,
            alignment=TA_RIGHT,
            spaceAfter=8,
        ),
        "body_en": ParagraphStyle(
            "MadarikBodyEnglish",
            parent=base["Normal"],
            fontName=ENGLISH_FONT_NAME,
            fontSize=9.5,
            leading=14,
            alignment=TA_LEFT,
            leftIndent=10,
            textColor=colors.HexColor("#334155"),
            spaceAfter=5,
        ),
        "question": ParagraphStyle(
            "MadarikQuestionTitle",
            parent=base["Normal"],
            fontName=font,
            fontSize=12,
            leading=18,
            alignment=TA_RIGHT,
            textColor=colors.HexColor("#0f172a"),
            spaceBefore=10,
            spaceAfter=4,
        ),
        "part": ParagraphStyle(
            "MadarikQuestionPartTitle",
            parent=base["Normal"],
            fontName=font,
            fontSize=10.5,
            leading=15,
            alignment=TA_RIGHT,
            rightIndent=8,
            textColor=colors.HexColor("#0f3b67"),
            spaceBefore=5,
            spaceAfter=2,
        ),
        "small": ParagraphStyle(
            "MadarikSmall",
            parent=base["Normal"],
            fontName=font,
            fontSize=8.5,
            leading=12,
            alignment=TA_RIGHT,
            textColor=colors.HexColor("#64748b"),
        ),
    }


def _add_pdf_logo(story: list, project: ProjectSession) -> None:
    logo_bytes = _logo_bytes(project)
    if not logo_bytes:
        return

    try:
        image = RLImage(BytesIO(logo_bytes))
        max_width = 2.2 * cm
        max_height = 2.2 * cm
        scale = min(max_width / image.drawWidth, max_height / image.drawHeight, 1)
        image.drawWidth *= scale
        image.drawHeight *= scale
        image.hAlign = "CENTER"
        story.append(image)
        story.append(Spacer(1, 0.15 * cm))
    except Exception:
        # Keep export resilient if the image payload is invalid.
        return


def _add_pdf_header(story: list, project: ProjectSession, questions: list[QuestionItem], styles: dict[str, ParagraphStyle]) -> None:
    metadata = project.metadata
    marks_total = sum(_question_total_marks(question) or 0 for question in questions)
    mode = "ثنائية اللغة" if metadata.output_mode == OutputMode.bilingual else "عربية نظيفة"

    story.append(_pdf_paragraph(metadata.paper_title or "ورقة تدريبية مترجمة", styles["title"], rtl=True))
    story.append(_pdf_paragraph("منصة مدارك", styles["subtitle"], rtl=True))

    rows = [
        ("اسم المدرسة", metadata.school_name, "المادة", metadata.subject),
        ("الصف", metadata.grade, "الفصل الدراسي", metadata.semester),
        ("الزمن", metadata.duration, "الدرجة", metadata.total_marks or str(marks_total)),
        ("اسم المعلم", metadata.teacher_name, "التاريخ", metadata.date),
        ("نوع النسخة", mode, "عدد الأسئلة", str(len(questions))),
    ]

    table_data = []
    for right_label, right_value, left_label, left_value in rows:
        table_data.append([
            _pdf_paragraph(f"{left_label}: {left_value or '__________'}", styles["small"], rtl=True),
            _pdf_paragraph(f"{right_label}: {right_value or '__________'}", styles["small"], rtl=True),
        ])

    table = Table(table_data, colWidths=[8.1 * cm, 8.1 * cm], hAlign="CENTER")
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 10))


def _add_pdf_question_assets(story: list, question: QuestionItem, styles: dict[str, ParagraphStyle]) -> None:
    if not question.attachments:
        return

    story.append(_pdf_paragraph("مرفقات السؤال:", styles["small"], rtl=True))
    for asset in question.attachments:
        asset_bytes = _asset_bytes(asset.data_base64)
        if not asset_bytes:
            continue
        try:
            image = RLImage(BytesIO(asset_bytes))
            max_width = 13.5 * cm
            max_height = 8.5 * cm
            scale = min(max_width / image.drawWidth, max_height / image.drawHeight, 1)
            image.drawWidth *= scale
            image.drawHeight *= scale
            image.hAlign = "CENTER"
            story.append(image)
            story.append(Spacer(1, 0.2 * cm))
        except Exception:
            story.append(_pdf_paragraph(f"تعذر إدراج المرفق: {asset.name}", styles["small"], rtl=True))


def _add_pdf_question_parts(
    story: list,
    project: ProjectSession,
    question: QuestionItem,
    styles: dict[str, ParagraphStyle],
) -> None:
    for part in _sorted_question_parts(question):
        depth = _question_part_depth(
            part,
            question.parts,
        )
        heading_style = styles["part"]

        if depth:
            heading_style = ParagraphStyle(
                f"{heading_style.name}Depth{depth}",
                parent=heading_style,
                rightIndent=(
                    heading_style.rightIndent
                    + (depth * 18)
                ),
            )

        story.append(
            _pdf_paragraph(
                _part_heading_text(part),
                heading_style,
                rtl=False,
            )
        )

        has_part_text = bool(
            part.original_text.strip()
            or part.translated_text.strip()
        )

        if not has_part_text:
            continue

        if project.metadata.output_mode == OutputMode.bilingual:
            _add_pdf_text_blocks(
                story,
                part.original_text,
                styles["body_en"],
                rtl=False,
                fallback="[Original text unavailable]",
                indent_level=depth,
            )
            _add_pdf_text_blocks(
                story,
                part.translated_text,
                styles["body_ar"],
                rtl=True,
                fallback="[تحتاج ترجمة]",
                indent_level=depth,
            )
        else:
            _add_pdf_text_blocks(
                story,
                part.translated_text or part.original_text,
                styles["body_ar"],
                rtl=True,
                fallback="[تحتاج ترجمة]",
                indent_level=depth,
            )




def _add_pdf_questions(
    story: list,
    project: ProjectSession,
    questions: list[QuestionItem],
    styles: dict[str, ParagraphStyle],
) -> None:
    for display_number, question in enumerate(
        questions,
        start=1,
    ):
        story.append(
            _pdf_paragraph(
                _question_heading_text(
                    display_number,
                    question,
                ),
                styles["question"],
                rtl=True,
            )
        )

        if question.parts:
            _add_pdf_question_parts(
                story,
                project,
                question,
                styles,
            )
        elif (
            project.metadata.output_mode
            == OutputMode.bilingual
        ):
            _add_pdf_text_blocks(
                story,
                question.original_text,
                styles["body_en"],
                rtl=False,
                fallback="[Original text unavailable]",
            )
            _add_pdf_text_blocks(
                story,
                question.translated_text,
                styles["body_ar"],
                rtl=True,
                fallback="[تحتاج ترجمة]",
            )
        else:
            _add_pdf_text_blocks(
                story,
                (
                    question.translated_text
                    or question.original_text
                ),
                styles["body_ar"],
                rtl=True,
                fallback="[تحتاج ترجمة]",
            )

        _add_pdf_question_assets(
            story,
            question,
            styles,
        )

        attachment_note = (
            _attachment_note_for_export(question)
        )

        if attachment_note:
            story.append(
                _pdf_paragraph(
                    f"ملاحظة مرفق: {attachment_note}",
                    styles["small"],
                    rtl=True,
                )
            )

def _draw_pdf_footer(canvas, doc) -> None:
    font = _ensure_pdf_font()
    canvas.saveState()
    canvas.setFont(font, 8)
    canvas.setFillColor(colors.HexColor("#64748b"))
    footer = _shape_arabic("أُنشئت هذه الورقة عبر منصة مدارك - مراجعة المعلم ضرورية قبل الاستخدام الصفي")
    canvas.drawCentredString(A4[0] / 2, 1.05 * cm, footer)
    canvas.drawRightString(A4[0] - 1.4 * cm, 1.05 * cm, str(doc.page))
    canvas.restoreState()


def build_project_pdf_bytes(project: ProjectSession) -> bytes:
    """Build a Phase 1-F2 RTL-friendly PDF for active project questions."""

    questions = _active_questions(project)
    if not questions:
        raise ValueError("لا توجد أسئلة نشطة قابلة للتصدير.")

    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=1.4 * cm,
        leftMargin=1.4 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.8 * cm,
        title=project.metadata.paper_title or "Madarik Export",
        author="Madarik Platform",
    )
    styles = _pdf_styles()
    story: list = []

    _add_pdf_logo(story, project)
    _add_pdf_header(story, project, questions, styles)
    _add_pdf_questions(story, project, questions, styles)

    document.build(story, onFirstPage=_draw_pdf_footer, onLaterPages=_draw_pdf_footer)
    return output.getvalue()


def safe_pdf_filename(project: ProjectSession) -> str:
    title = project.metadata.paper_title or "madarik_export"
    subject = project.metadata.subject or "paper"
    raw = f"madarik_{subject}_{title}"
    normalized = re.sub(r"[^A-Za-z0-9_\-]+", "_", raw).strip("_")
    if not normalized:
        normalized = "madarik_export"
    return f"{normalized[:90]}.pdf"
