from __future__ import annotations

from pathlib import Path


from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.services.file_names import safe_filename_stem
from app.models.assessment import (
    AssessmentDraft,
    AssessmentExportResponse,
    AssessmentStudentPaperPreview,
    AssessmentStudentPaperQuestion,
    AssessmentStudentPaperSection,
)
from app.services.assessment_builder import (
    build_assessment_detail,
)
from app.services.question_bank_repository import (
    QuestionBankRepository,
)


EXPORT_DIR = Path("data/exports/assessments")


def build_student_paper_preview(
    draft: AssessmentDraft,
    question_bank_repository: QuestionBankRepository,
) -> AssessmentStudentPaperPreview:
    detail = build_assessment_detail(
        draft,
        question_bank_repository,
    )

    section_map = {
        section.id: AssessmentStudentPaperSection(
            id=section.id,
            title=section.title,
            instructions=section.instructions,
            order_index=section.order_index,
        )
        for section in sorted(
            draft.sections,
            key=lambda section: section.order_index,
        )
    }

    fallback_section = AssessmentStudentPaperSection(
        id="unsectioned",
        title="أسئلة دون قسم",
        order_index=9999,
    )

    for number, question in enumerate(
        sorted(
            detail.questions,
            key=lambda item: item.order_index,
        ),
        start=1,
    ):
        target = (
            section_map.get(question.section_id)
            if question.section_id
            else None
        ) or fallback_section

        target.questions.append(
            AssessmentStudentPaperQuestion(
                bank_item_id=question.bank_item_id,
                number=number,
                question_number=question.question_number,
                text=question.text,
                marks=question.marks,
                section_id=question.section_id,
                section_title=target.title,
            )
        )

    sections = [
        section
        for section in sorted(
            section_map.values(),
            key=lambda section: section.order_index,
        )
        if section.questions
    ]
    if fallback_section.questions:
        sections.append(fallback_section)

    issues: list[str] = []
    if not draft.blueprint.title.strip():
        issues.append("عنوان الاختبار غير موجود.")
    if not sections:
        issues.append("لا توجد أسئلة في المسودة.")
    if detail.balance.selected_marks != draft.blueprint.total_marks:
        issues.append(
            "مجموع درجات الأسئلة لا يطابق الدرجة الكلية."
        )
    if len(detail.questions) != draft.blueprint.target_question_count:
        issues.append(
            "عدد الأسئلة الفعلي لا يطابق العدد المستهدف."
        )

    return AssessmentStudentPaperPreview(
        draft_id=draft.id,
        title=draft.blueprint.title,
        grade=draft.blueprint.grade,
        science_domain=draft.blueprint.science_domain,
        subject_id=draft.blueprint.subject_id,
        duration_minutes=draft.blueprint.duration_minutes,
        total_marks=draft.blueprint.total_marks,
        question_count=len(detail.questions),
        sections=sections,
        export_ready=not issues,
        issues=issues,
    )


def _render_plain_text(
    preview: AssessmentStudentPaperPreview,
) -> str:
    lines = [
        preview.title,
        f"الصف: {preview.grade}",
        f"الزمن: {preview.duration_minutes} دقيقة",
        f"الدرجة الكلية: {preview.total_marks}",
        "",
    ]

    for section in preview.sections:
        lines.append(section.title)
        if section.instructions:
            lines.append(section.instructions)
        lines.append("")

        for question in section.questions:
            lines.append(
                f"{question.number}. "
                f"{question.text} "
                f"({question.marks} درجات)"
            )
            lines.append("")

    lines.append("صفحة الإجابة")
    lines.append("")
    for section in preview.sections:
        for question in section.questions:
            lines.append(
                f"{question.number}. "
                + "_" * 60
            )

    return "\n".join(lines)



def _set_paragraph_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(4)

    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def _set_cell_rtl(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    bidi = OxmlElement("w:bidiVisual")
    bidi.set(qn("w:val"), "1")
    tc_pr.append(bidi)

    for paragraph in cell.paragraphs:
        _set_paragraph_rtl(paragraph)


def _add_docx_header(
    document: Document,
    preview: AssessmentStudentPaperPreview,
) -> None:
    title = document.add_paragraph()
    _set_paragraph_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(preview.title)
    run.bold = True
    run.font.size = Pt(18)

    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"

    values = [
        ("الصف", str(preview.grade)),
        ("الزمن", f"{preview.duration_minutes} دقيقة"),
        ("الدرجة", str(preview.total_marks)),
    ]

    for index, (label, value) in enumerate(values):
        cell = table.cell(0, index)
        cell.text = label
        _set_cell_rtl(cell)
        cell.paragraphs[0].runs[0].bold = True

        value_cell = table.cell(1, index)
        value_cell.text = value
        _set_cell_rtl(value_cell)

    document.add_paragraph()


def _add_docx_question(
    document: Document,
    question: AssessmentStudentPaperQuestion,
) -> None:
    paragraph = document.add_paragraph()
    _set_paragraph_rtl(paragraph)

    number_run = paragraph.add_run(f"{question.number}. ")
    number_run.bold = True

    text_run = paragraph.add_run(question.text)
    text_run.font.size = Pt(12)

    marks_run = paragraph.add_run(
        f"  ({question.marks} درجات)"
    )
    marks_run.bold = True

    for _ in range(2):
        line = document.add_paragraph(
            "................................................................................"
        )
        _set_paragraph_rtl(line)


def _build_docx(
    preview: AssessmentStudentPaperPreview,
    path: Path,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _add_docx_header(document, preview)

    for assessment_section in preview.sections:
        heading = document.add_paragraph()
        _set_paragraph_rtl(heading)
        run = heading.add_run(assessment_section.title)
        run.bold = True
        run.font.size = Pt(14)

        if assessment_section.instructions:
            instructions = document.add_paragraph(
                assessment_section.instructions
            )
            _set_paragraph_rtl(instructions)

        for question in assessment_section.questions:
            _add_docx_question(document, question)

    document.add_section(WD_SECTION.NEW_PAGE)

    answer_heading = document.add_paragraph()
    _set_paragraph_rtl(answer_heading)
    answer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = answer_heading.add_run("صفحة الإجابة")
    run.bold = True
    run.font.size = Pt(16)

    for assessment_section in preview.sections:
        for question in assessment_section.questions:
            line = document.add_paragraph(
                f"{question.number}. "
                + "_" * 70
            )
            _set_paragraph_rtl(line)

    document.save(path)


def _register_pdf_font() -> str:
    candidates = (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
    )
    for candidate in candidates:
        if candidate.exists():
            pdfmetrics.registerFont(
                TTFont("AssessmentArabic", str(candidate))
            )
            return "AssessmentArabic"
    return "Helvetica"


def _build_pdf(
    preview: AssessmentStudentPaperPreview,
    path: Path,
) -> None:
    font_name = _register_pdf_font()
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "AssessmentTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=18,
        leading=24,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "AssessmentHeading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=14,
        leading=20,
        alignment=TA_RIGHT,
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "AssessmentBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=11,
        leading=18,
        alignment=TA_RIGHT,
    )

    story = [
        Paragraph(preview.title, title_style),
        Table(
            [
                ["الصف", "الزمن", "الدرجة"],
                [
                    str(preview.grade),
                    f"{preview.duration_minutes} دقيقة",
                    str(preview.total_marks),
                ],
            ],
            colWidths=[5.5 * cm, 5.5 * cm, 5.5 * cm],
            style=TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.5, "black"),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                ]
            ),
        ),
        Spacer(1, 0.5 * cm),
    ]

    for assessment_section in preview.sections:
        story.append(
            Paragraph(
                assessment_section.title,
                heading_style,
            )
        )
        if assessment_section.instructions:
            story.append(
                Paragraph(
                    assessment_section.instructions,
                    body_style,
                )
            )

        for question in assessment_section.questions:
            story.append(
                Paragraph(
                    f"{question.number}. {question.text} "
                    f"({question.marks} درجات)",
                    body_style,
                )
            )
            story.append(
                Paragraph(
                    "." * 90,
                    body_style,
                )
            )
            story.append(
                Paragraph(
                    "." * 90,
                    body_style,
                )
            )
            story.append(Spacer(1, 0.2 * cm))

    story.append(PageBreak())
    story.append(
        Paragraph("صفحة الإجابة", title_style)
    )

    for assessment_section in preview.sections:
        for question in assessment_section.questions:
            story.append(
                Paragraph(
                    f"{question.number}. " + "_" * 80,
                    body_style,
                )
            )
            story.append(Spacer(1, 0.25 * cm))

    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )
    document.build(story)


def export_assessment_foundation(
    draft: AssessmentDraft,
    question_bank_repository: QuestionBankRepository,
    output_format: str,
) -> AssessmentExportResponse:
    preview = build_student_paper_preview(
        draft,
        question_bank_repository,
    )

    if not preview.export_ready:
        return AssessmentExportResponse(
            draft_id=draft.id,
            format=output_format,
            filename="",
            path="",
            export_ready=False,
            issues=preview.issues,
        )

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    safe_title = safe_filename_stem(
        draft.blueprint.title,
        fallback="assessment",
    )

    if output_format == "docx":
        filename = f"{safe_title}-{draft.id}.docx"
    else:
        filename = f"{safe_title}-{draft.id}.pdf"

    path = EXPORT_DIR / filename

    if output_format == "docx":
        _build_docx(preview, path)
    else:
        _build_pdf(preview, path)

    return AssessmentExportResponse(
        draft_id=draft.id,
        format=output_format,
        filename=filename,
        path=str(path),
        export_ready=True,
        issues=[],
    )
