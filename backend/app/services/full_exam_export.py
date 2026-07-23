from __future__ import annotations

import base64
import hashlib
from io import BytesIO
import re
from typing import Iterable
from zipfile import BadZipFile, ZipFile

from app.services.arabic_text import reshape_arabic
from docx import Document
from PIL import Image as PILImage
from pypdf import PdfReader

from app.models.project import (
    ExportFormat,
    FullExamExportAcceptanceStatus,
    FullExamExportArtifactStatus,
    FullExamExportCheck,
    FullExamExportFormatSummary,
    FullExamExportReport,
    FullExamIntakeStatus,
    FullExamTranslationAcceptanceStatus,
    ProjectSession,
    QuestionItem,
    QuestionPart,
    QuestionStatus,
)
from app.services.export_review import (
    marks_policy_resolves_total,
    missing_visual_asset_question_numbers,
    parse_declared_total_marks,
)


EXPORT_MANIFEST_PREFIX = "MADARIK-A6C-V1"
_MANIFEST_FIELD_RE = re.compile(r"(?P<key>[a-z_]+)=(?P<value>[^;]*)")
_DOCX_QUESTION_HEADING_RE = re.compile(
    r"^السؤال\s+(?P<number>\d+)(?:\s+\[(?P<marks>\d+)\])?$"
)


def _active_questions(project: ProjectSession) -> list[QuestionItem]:
    return sorted(
        (
            question
            for question in project.questions
            if question.status != QuestionStatus.deleted
        ),
        key=lambda question: question.order_index,
    )


_PDF_BIDI_CONTROL_RE = re.compile(
    r"[\u200e\u200f\u202a-\u202e\u2066-\u2069]"
)
_PDF_DIGIT_TRANSLATION = str.maketrans(
    "٠١٢٣٤٥٦٧٨٩۰۱۲۳۴۵۶۷۸۹",
    "01234567890123456789",
)


def _normalize_pdf_extracted_text(value: str) -> str:
    normalized = value.translate(_PDF_DIGIT_TRANSLATION)
    normalized = _PDF_BIDI_CONTROL_RE.sub("", normalized)
    return normalized.replace("\u00a0", " ")


def _pdf_heading_signatures(
    extracted_text: str,
) -> set[tuple[int, int]]:
    """Extract stable ``question number + marks`` tokens from PDF text.

    ReportLab/PDF readers may join, split, or reverse the RTL heading fragments.
    The visible Arabic word is therefore deliberately ignored. The numeric
    heading signature remains stable and does not collide with a body-only
    marks token such as ``[2]``.
    """

    normalized = _normalize_pdf_extracted_text(extracted_text)
    lines = [
        re.sub(r"\s+", " ", line).strip()
        for line in normalized.splitlines()
        if line.strip()
    ]
    segments = list(lines)
    segments.extend(
        f"{lines[index]} {lines[index + 1]}"
        for index in range(len(lines) - 1)
    )

    signatures: set[tuple[int, int]] = set()
    for segment in segments:
        compact = re.sub(r"\s+", "", segment)

        for match in re.finditer(
            r"(?<!\d)(?P<number>\d+)\[(?P<marks>\d+)\](?!\d)",
            compact,
        ):
            signatures.add(
                (
                    int(match.group("number")),
                    int(match.group("marks")),
                )
            )

        for match in re.finditer(
            r"(?<!\d)\[(?P<marks>\d+)\](?P<number>\d+)(?!\d)",
            compact,
        ):
            signatures.add(
                (
                    int(match.group("number")),
                    int(match.group("marks")),
                )
            )

    return signatures


def _pdf_question_heading_count(
    project: ProjectSession,
    extracted_text: str,
) -> int:
    """Count rendered PDF question headings across RTL extractor variants."""

    normalized = _normalize_pdf_extracted_text(extracted_text)
    exact_lines = {
        re.sub(r"\s+", " ", line).strip()
        for line in normalized.splitlines()
        if line.strip()
    }
    signatures = _pdf_heading_signatures(normalized)
    detected = 0

    for display_number, question in enumerate(
        _active_questions(project),
        start=1,
    ):
        marks = question_export_total_marks(question)

        if marks is not None:
            if (display_number, marks) in signatures:
                detected += 1
            continue

        # Unmarked headings have no safe number+marks signature. Preserve the
        # conservative exact-line fallback used before this fix.
        if str(display_number) in exact_lines:
            detected += 1

    return detected


def _question_part_depth(
    part: QuestionPart,
    parts: list[QuestionPart],
) -> int:
    parts_by_id = {item.id: item for item in parts}
    current = part
    visited = {part.id}
    depth = 0

    while current.parent_id:
        parent = parts_by_id.get(current.parent_id)
        if parent is None or parent.id in visited:
            break
        visited.add(parent.id)
        current = parent
        depth += 1

    return depth


def question_export_total_marks(question: QuestionItem) -> int | None:
    """Return the same hierarchy-safe total used by the export renderers."""

    if question.marks is not None:
        return question.marks

    if not question.parts:
        return None

    parts_by_id = {part.id: part for part in question.parts}
    children: dict[str, list[QuestionPart]] = {}

    for part in question.parts:
        if part.parent_id in parts_by_id:
            children.setdefault(part.parent_id, []).append(part)

    roots = [
        part
        for part in question.parts
        if part.parent_id not in parts_by_id
    ]

    def branch_marks(
        part: QuestionPart,
        visited: set[str],
    ) -> int | None:
        if part.id in visited:
            return part.marks

        child_marks = [
            branch_marks(child, visited | {part.id})
            for child in children.get(part.id, [])
        ]
        concrete_child_marks = [
            value
            for value in child_marks
            if value is not None
        ]

        if concrete_child_marks:
            return sum(concrete_child_marks)

        return part.marks

    root_totals = [branch_marks(root, set()) for root in roots]
    concrete_totals = [
        value
        for value in root_totals
        if value is not None
    ]
    return sum(concrete_totals) if concrete_totals else None


def _sequence_digest(values: Iterable[str]) -> str:
    source = "\x1f".join(values)
    return hashlib.sha256(source.encode("utf-8")).hexdigest()[:20]


def _expected_structure(project: ProjectSession) -> dict[str, object]:
    questions = _active_questions(project)
    return {
        "questions": len(questions),
        "parts": sum(len(question.parts) for question in questions),
        "attachments": sum(len(question.attachments) for question in questions),
        "marks": sum(question_export_total_marks(question) or 0 for question in questions),
        "order": ",".join(str(index) for index in range(1, len(questions) + 1)),
        "sequence": _sequence_digest(question.id for question in questions),
    }


def build_full_exam_export_manifest(project: ProjectSession) -> str:
    """Build the backwards-compatible expected-structure artifact manifest."""

    structure = _expected_structure(project)
    return ";".join(
        [
            EXPORT_MANIFEST_PREFIX,
            f"questions={structure['questions']}",
            f"parts={structure['parts']}",
            f"attachments={structure['attachments']}",
            f"marks={structure['marks']}",
            f"order={structure['order']}",
            f"sequence={structure['sequence']}",
        ]
    )


def build_full_exam_pdf_render_manifest(
    project: ProjectSession,
    *,
    rendered_question_ids: Iterable[str],
    rendered_part_ids: Iterable[str],
    rendered_attachment_ids: Iterable[str],
    rendered_marks: int,
    rendered_order: Iterable[int],
) -> str:
    """Build PDF metadata from flowables successfully queued for rendering.

    The expected structure remains in the V1-compatible fields. The additional
    render-evidence fields are created from the actual question, part, and image
    flowables appended to the PDF story before ``document.build`` succeeds.
    """

    question_ids = list(rendered_question_ids)
    part_ids = list(rendered_part_ids)
    attachment_ids = list(rendered_attachment_ids)
    order = list(rendered_order)

    return ";".join(
        [
            build_full_exam_export_manifest(project),
            "render_evidence=1",
            f"rendered_questions={len(question_ids)}",
            f"rendered_parts={len(part_ids)}",
            f"rendered_attachments={len(attachment_ids)}",
            f"rendered_marks={rendered_marks}",
            f"rendered_order={','.join(str(item) for item in order)}",
            f"rendered_sequence={_sequence_digest(question_ids)}",
            f"rendered_attachment_sequence={_sequence_digest(attachment_ids)}",
        ]
    )


def _parse_manifest(value: str | None) -> dict[str, object] | None:
    manifest = (value or "").strip()
    if not manifest.startswith(f"{EXPORT_MANIFEST_PREFIX};"):
        return None

    fields = {
        match.group("key"): match.group("value")
        for match in _MANIFEST_FIELD_RE.finditer(manifest)
    }

    try:
        result: dict[str, object] = {
            "questions": int(fields["questions"]),
            "parts": int(fields["parts"]),
            "attachments": int(fields["attachments"]),
            "marks": int(fields["marks"]),
            "order": fields.get("order", ""),
            "sequence": fields.get("sequence", ""),
            "render_evidence": fields.get("render_evidence") == "1",
        }
    except (KeyError, TypeError, ValueError):
        return None

    if not result["render_evidence"]:
        return result

    try:
        result.update(
            {
                "rendered_questions": int(fields["rendered_questions"]),
                "rendered_parts": int(fields["rendered_parts"]),
                "rendered_attachments": int(fields["rendered_attachments"]),
                "rendered_marks": int(fields["rendered_marks"]),
                "rendered_order": fields.get("rendered_order", ""),
                "rendered_sequence": fields["rendered_sequence"],
                "rendered_attachment_sequence": fields[
                    "rendered_attachment_sequence"
                ],
                "render_evidence_valid": True,
            }
        )
    except (KeyError, TypeError, ValueError):
        result["render_evidence_valid"] = False

    return result


def _is_valid_image_payload(data_base64: str) -> bool:
    if not data_base64:
        return False
    try:
        payload = base64.b64decode(data_base64, validate=True)
        with PILImage.open(BytesIO(payload)) as image:
            image.verify()
        return True
    except Exception:
        return False


def _valid_logo_count(project: ProjectSession) -> int:
    if not project.school_logo:
        return 0
    return int(_is_valid_image_payload(project.school_logo.data_base64))


def _valid_attachment_ids_in_order(project: ProjectSession) -> list[str]:
    return [
        attachment.id
        for question in _active_questions(project)
        for attachment in question.attachments
        if _is_valid_image_payload(attachment.data_base64)
    ]


def _valid_attachment_ids(project: ProjectSession) -> set[str]:
    return set(_valid_attachment_ids_in_order(project))


def _valid_attachment_count(project: ProjectSession) -> int:
    return len(_valid_attachment_ids_in_order(project))


def _docx_text(document: Document) -> str:
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                blocks.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(blocks)


def _docx_question_headings(document: Document) -> tuple[list[int], int]:
    order: list[int] = []
    marks_total = 0
    for paragraph in document.paragraphs:
        match = _DOCX_QUESTION_HEADING_RE.match(paragraph.text.strip())
        if not match:
            continue
        order.append(int(match.group("number")))
        if match.group("marks"):
            marks_total += int(match.group("marks"))
    return order, marks_total


def _resolve_pdf_object(value: object) -> object:
    try:
        return value.get_object()  # type: ignore[attr-defined]
    except Exception:
        return value


def _pdf_object_identity(
    reference: object,
    object_value: object,
) -> tuple[str, int, int] | tuple[str, int]:
    if hasattr(reference, "idnum"):
        return (
            "reference",
            int(reference.idnum),  # type: ignore[attr-defined]
            int(getattr(reference, "generation", 0)),
        )

    indirect_reference = getattr(
        object_value,
        "indirect_reference",
        None,
    )
    if indirect_reference is not None and hasattr(
        indirect_reference,
        "idnum",
    ):
        return (
            "reference",
            int(indirect_reference.idnum),
            int(getattr(indirect_reference, "generation", 0)),
        )

    return ("object", id(object_value))


def _pdf_image_count(reader: PdfReader) -> int:
    """Count visible PDF images across direct and nested Form XObjects.

    PDF producers are free to place an image directly in page resources or
    inside one or more Form XObjects. Alpha soft masks are also represented as
    image XObjects but must not be counted as a separate exported attachment.
    """

    visible_images: set[
        tuple[str, int, int] | tuple[str, int]
    ] = set()
    mask_images: set[
        tuple[str, int, int] | tuple[str, int]
    ] = set()
    visited_forms: set[
        tuple[str, int, int] | tuple[str, int]
    ] = set()

    def visit_resources(resources_value: object | None) -> None:
        if resources_value is None:
            return

        resources = _resolve_pdf_object(resources_value)
        if not hasattr(resources, "get"):
            return

        xobjects_value = resources.get("/XObject")  # type: ignore[attr-defined]
        if xobjects_value is None:
            return

        xobjects = _resolve_pdf_object(xobjects_value)
        if not hasattr(xobjects, "items"):
            return

        for _, reference in xobjects.items():  # type: ignore[attr-defined]
            object_value = _resolve_pdf_object(reference)
            if not hasattr(object_value, "get"):
                continue

            identity = _pdf_object_identity(
                reference,
                object_value,
            )
            subtype = object_value.get("/Subtype")  # type: ignore[attr-defined]

            if subtype == "/Image":
                if bool(
                    object_value.get("/ImageMask")  # type: ignore[attr-defined]
                ):
                    mask_images.add(identity)
                else:
                    visible_images.add(identity)

                soft_mask_reference = object_value.get(  # type: ignore[attr-defined]
                    "/SMask"
                )
                if soft_mask_reference is not None:
                    soft_mask = _resolve_pdf_object(
                        soft_mask_reference
                    )
                    mask_images.add(
                        _pdf_object_identity(
                            soft_mask_reference,
                            soft_mask,
                        )
                    )
                continue

            if subtype != "/Form" or identity in visited_forms:
                continue

            visited_forms.add(identity)
            visit_resources(
                object_value.get("/Resources")  # type: ignore[attr-defined]
            )

    for page in reader.pages:
        visit_resources(page.get("/Resources"))

    return len(visible_images - mask_images)


def _metadata_value(metadata: object, key: str) -> str:
    if metadata is None:
        return ""
    value = getattr(metadata, key, None)
    if value is None and isinstance(metadata, dict):
        value = metadata.get(key)
    return str(value or "")


def _manifest_checks(
    project: ProjectSession,
    manifest: dict[str, object] | None,
) -> list[FullExamExportCheck]:
    expected = _expected_structure(project)

    if manifest is None:
        return [
            FullExamExportCheck(
                code="artifact_manifest",
                passed=False,
                message="لم يُعثر على بصمة بنية A6c داخل ملف التصدير.",
            )
        ]

    fields = (
        "questions",
        "parts",
        "attachments",
        "marks",
        "order",
        "sequence",
    )
    return [
        FullExamExportCheck(
            code=f"manifest_{field}",
            passed=manifest[field] == expected[field],
            message=(
                f"تطابق حقل {field} في بصمة ملف التصدير."
                if manifest[field] == expected[field]
                else (
                    f"حقل {field} في الملف ({manifest[field]}) لا يطابق "
                    f"بنية المشروع الحالية ({expected[field]})."
                )
            ),
        )
        for field in fields
    ]


def _pdf_render_evidence(
    project: ProjectSession,
    manifest: dict[str, object] | None,
) -> tuple[bool, list[FullExamExportCheck]] | None:
    """Validate deterministic PDF render evidence when the artifact has it.

    Old artifacts without render evidence keep the extractor-based fallback.
    New artifacts fail closed when the evidence marker exists but is incomplete
    or differs from the active project structure.
    """

    if manifest is None or not manifest.get("render_evidence"):
        return None

    expected = _expected_structure(project)
    valid_attachment_ids = _valid_attachment_ids_in_order(project)
    expected_attachment_sequence = _sequence_digest(valid_attachment_ids)

    if not manifest.get("render_evidence_valid"):
        return (
            False,
            [
                FullExamExportCheck(
                    code="pdf_render_evidence",
                    passed=False,
                    message="بصمة دليل توليد PDF موجودة لكنها ناقصة أو غير صالحة.",
                )
            ],
        )

    comparisons = [
        (
            "rendered_questions",
            expected["questions"],
            "عدد الأسئلة التي أضيفت فعليًا إلى قصة PDF",
        ),
        (
            "rendered_parts",
            expected["parts"],
            "عدد أجزاء الأسئلة التي أضيفت فعليًا إلى قصة PDF",
        ),
        (
            "rendered_attachments",
            len(valid_attachment_ids),
            "عدد مرفقات الأسئلة الصالحة التي أضيفت فعليًا إلى قصة PDF",
        ),
        (
            "rendered_marks",
            expected["marks"],
            "مجموع الدرجات في عناوين الأسئلة المولدة داخل PDF",
        ),
        (
            "rendered_order",
            expected["order"],
            "ترتيب الأسئلة المولد داخل PDF",
        ),
        (
            "rendered_sequence",
            expected["sequence"],
            "تسلسل هويات الأسئلة المولدة داخل PDF",
        ),
        (
            "rendered_attachment_sequence",
            expected_attachment_sequence,
            "تسلسل هويات المرفقات المولدة داخل PDF",
        ),
    ]

    checks = [
        FullExamExportCheck(
            code=f"pdf_{field}",
            passed=manifest.get(field) == expected_value,
            message=(
                f"تطابق {label}."
                if manifest.get(field) == expected_value
                else (
                    f"لا يطابق {label}: القيمة المولدة "
                    f"({manifest.get(field)}) والمتوقعة ({expected_value})."
                )
            ),
        )
        for field, expected_value, label in comparisons
    ]
    valid = all(check.passed for check in checks)
    checks.insert(
        0,
        FullExamExportCheck(
            code="pdf_render_evidence",
            passed=valid,
            message=(
                "دليل توليد PDF مكتمل ومتوافق مع بنية المشروع."
                if valid
                else "دليل توليد PDF لا يطابق بنية المشروع الحالية."
            ),
        ),
    )
    return valid, checks


def _inspect_docx(
    project: ProjectSession,
    artifact_bytes: bytes,
) -> FullExamExportFormatSummary:
    checks: list[FullExamExportCheck] = []
    warnings: list[str] = []
    manifest: dict[str, object] | None = None
    question_order: list[int] = []
    heading_marks = 0
    exported_images = 0
    body_text = ""

    try:
        with ZipFile(BytesIO(artifact_bytes)) as archive:
            checks.append(
                FullExamExportCheck(
                    code="docx_container",
                    passed="word/document.xml" in archive.namelist(),
                    message="ملف Word حاوية DOCX صالحة وقابلة للقراءة.",
                )
            )
    except BadZipFile:
        checks.append(
            FullExamExportCheck(
                code="docx_container",
                passed=False,
                message="ملف Word الناتج ليس حاوية DOCX صالحة.",
            )
        )

    try:
        document = Document(BytesIO(artifact_bytes))
        manifest = _parse_manifest(document.core_properties.subject)
        question_order, heading_marks = _docx_question_headings(document)
        exported_images = max(
            0,
            len(document.inline_shapes) - _valid_logo_count(project),
        )
        body_text = _docx_text(document)
    except Exception as exc:
        warnings.append(f"تعذر فحص محتوى Word تفصيليًا: {exc}")

    checks.extend(_manifest_checks(project, manifest))
    expected = _expected_structure(project)
    expected_order = list(range(1, int(expected["questions"]) + 1))
    valid_attachment_ids = _valid_attachment_ids(project)
    valid_attachments = len(valid_attachment_ids)
    missing_visual_questions = missing_visual_asset_question_numbers(
        _active_questions(project),
        valid_attachment_ids=valid_attachment_ids,
    )

    checks.extend(
        [
            FullExamExportCheck(
                code="docx_question_order",
                passed=question_order == expected_order,
                message=(
                    "ترتيب عناوين الأسئلة داخل Word مطابق للترتيب النشط."
                    if question_order == expected_order
                    else (
                        f"ترتيب عناوين Word المكتشف {question_order} لا يطابق "
                        f"الترتيب المتوقع {expected_order}."
                    )
                ),
            ),
            FullExamExportCheck(
                code="docx_marks_total",
                passed=heading_marks == expected["marks"],
                message=(
                    "مجموع درجات عناوين Word مطابق دون جمع درجات الأجزاء مرتين."
                    if heading_marks == expected["marks"]
                    else (
                        f"مجموع الدرجات المكتشف في عناوين Word ({heading_marks}) "
                        f"لا يطابق المتوقع ({expected['marks']})."
                    )
                ),
            ),
            FullExamExportCheck(
                code="docx_attachments_once",
                passed=exported_images == valid_attachments,
                message=(
                    "أُدرجت مرفقات الأسئلة الصالحة مرة واحدة في Word."
                    if exported_images == valid_attachments
                    else (
                        f"عدد صور الأسئلة داخل Word ({exported_images}) لا يطابق "
                        f"المرفقات الصالحة ({valid_attachments})."
                    )
                ),
            ),
            FullExamExportCheck(
                code="docx_visual_question_coverage",
                passed=not missing_visual_questions,
                message=(
                    "كل سؤال يعتمد على رسم أو مخطط يملك مرفقًا بصريًا صالحًا في Word."
                    if not missing_visual_questions
                    else (
                        "توجد أسئلة تعتمد على رسم أو مخطط دون مرفق بصري صالح في Word: "
                        + "، ".join(missing_visual_questions)
                    )
                ),
            ),
            FullExamExportCheck(
                code="docx_body_content",
                passed=bool(body_text.strip()),
                message="يحتوي ملف Word على نص قابل للقراءة.",
            ),
        ]
    )

    invalid_attachments = int(expected["attachments"]) - valid_attachments
    if invalid_attachments:
        warnings.append(
            f"تعذر إدراج {invalid_attachments} مرفق بسبب بيانات صورة غير صالحة."
        )

    status = (
        FullExamExportArtifactStatus.accepted
        if all(check.passed for check in checks) and not warnings
        else FullExamExportArtifactStatus.needs_review
    )
    if not artifact_bytes or not checks[0].passed or manifest is None:
        status = FullExamExportArtifactStatus.failed

    return FullExamExportFormatSummary(
        format=ExportFormat.docx,
        status=status,
        byte_size=len(artifact_bytes),
        page_count=None,
        exported_question_count=(
            int(manifest["questions"])
            if manifest is not None
            else len(question_order)
        ),
        exported_part_count=(
            int(manifest["parts"])
            if manifest is not None
            else 0
        ),
        exported_attachment_count=exported_images,
        detected_total_marks=(
            int(manifest["marks"])
            if manifest is not None
            else heading_marks
        ),
        question_order=[str(number) for number in question_order],
        checks=checks,
        warnings=warnings,
    )


def _inspect_pdf(
    project: ProjectSession,
    artifact_bytes: bytes,
) -> FullExamExportFormatSummary:
    checks: list[FullExamExportCheck] = []
    warnings: list[str] = []
    diagnostic_errors: list[str] = []
    manifest: dict[str, object] | None = None
    page_count = 0
    extracted_image_count = 0
    extracted_heading_count = 0
    extracted_text = ""

    signature_valid = artifact_bytes.startswith(b"%PDF")
    checks.append(
        FullExamExportCheck(
            code="pdf_signature",
            passed=signature_valid,
            message=(
                "ملف PDF يبدأ بتوقيع PDF صالح."
                if signature_valid
                else "ملف PDF الناتج لا يبدأ بتوقيع PDF صالح."
            ),
        )
    )

    reader: PdfReader | None = None
    try:
        reader = PdfReader(BytesIO(artifact_bytes))
        page_count = len(reader.pages)
        manifest = _parse_manifest(
            _metadata_value(reader.metadata, "subject")
        )
    except Exception as exc:
        warnings.append(f"تعذر فتح بنية PDF أو قراءة بياناته الوصفية: {exc}")

    if reader is not None:
        try:
            extracted_image_count = max(
                0,
                _pdf_image_count(reader) - _valid_logo_count(project),
            )
        except Exception as exc:
            diagnostic_errors.append(
                f"تعذر تشخيص XObject داخل PDF: {exc}"
            )

        try:
            extracted_text = "\n".join(
                page.extract_text() or ""
                for page in reader.pages
            )
            extracted_heading_count = _pdf_question_heading_count(
                project,
                extracted_text,
            )
        except Exception as exc:
            diagnostic_errors.append(
                f"تعذر تشخيص طبقة النص داخل PDF: {exc}"
            )

    checks.extend(_manifest_checks(project, manifest))
    expected = _expected_structure(project)
    valid_attachment_ids = _valid_attachment_ids(project)
    valid_attachments = len(valid_attachment_ids)
    missing_visual_questions = missing_visual_asset_question_numbers(
        _active_questions(project),
        valid_attachment_ids=valid_attachment_ids,
    )

    render_evidence = _pdf_render_evidence(project, manifest)
    if render_evidence is not None:
        render_evidence_valid, render_checks = render_evidence
        checks.extend(render_checks)
        rendered_questions = (
            int(manifest.get("rendered_questions", 0))
            if manifest
            else 0
        )
        rendered_parts = (
            int(manifest.get("rendered_parts", 0))
            if manifest
            else 0
        )
        rendered_attachments = (
            int(manifest.get("rendered_attachments", 0))
            if manifest
            else 0
        )
        rendered_marks = (
            int(manifest.get("rendered_marks", 0))
            if manifest
            else 0
        )
        rendered_order = (
            str(manifest.get("rendered_order", ""))
            if manifest
            else ""
        )

        heading_passed = (
            render_evidence_valid
            and rendered_questions == expected["questions"]
        )
        attachment_passed = (
            render_evidence_valid
            and rendered_attachments == valid_attachments
        )
        text_layer_passed = (
            render_evidence_valid
            and rendered_questions > 0
        )
        exported_question_count = rendered_questions
        exported_part_count = rendered_parts
        exported_attachment_count = rendered_attachments
        detected_total_marks = rendered_marks
        question_order = [
            item
            for item in rendered_order.split(",")
            if item
        ]
        heading_message = (
            "دليل التوليد يثبت إدراج جميع عناوين الأسئلة في PDF "
            f"(تشخيص قارئ النص اكتشف {extracted_heading_count})."
            if heading_passed
            else "دليل التوليد لا يثبت إدراج جميع عناوين الأسئلة في PDF."
        )
        attachment_message = (
            "دليل التوليد يثبت إدراج مرفقات الأسئلة الصالحة مرة واحدة "
            f"(تشخيص XObject اكتشف {extracted_image_count})."
            if attachment_passed
            else "دليل التوليد لا يثبت إدراج مرفقات الأسئلة الصالحة مرة واحدة."
        )
        text_layer_message = (
            "دليل التوليد يثبت إضافة عناصر نصية إلى PDF "
            f"(استخراج القارئ أعاد {len(extracted_text.strip())} محرفًا)."
            if text_layer_passed
            else "دليل التوليد لا يثبت إضافة عناصر نصية إلى PDF."
        )
    else:
        # Backwards-compatible fallback for PDF files generated before Fix 2.
        warnings.extend(diagnostic_errors)
        heading_passed = extracted_heading_count == expected["questions"]
        attachment_passed = extracted_image_count == valid_attachments
        text_layer_passed = bool(extracted_text.strip())
        exported_question_count = (
            int(manifest["questions"])
            if manifest
            else 0
        )
        exported_part_count = (
            int(manifest["parts"])
            if manifest
            else 0
        )
        exported_attachment_count = extracted_image_count
        detected_total_marks = (
            int(manifest["marks"])
            if manifest
            else 0
        )
        manifest_order = (
            str(manifest.get("order", ""))
            if manifest
            else ""
        )
        question_order = [
            item
            for item in manifest_order.split(",")
            if item
        ]
        heading_message = (
            "عدد عناوين الأسئلة المرئية داخل PDF مطابق للبنية المتوقعة."
            if heading_passed
            else (
                "عدد عناوين الأسئلة المكتشف داخل PDF لا يطابق "
                "عدد الأسئلة النشطة."
            )
        )
        attachment_message = (
            "أُدرجت مرفقات الأسئلة الصالحة مرة واحدة في PDF."
            if attachment_passed
            else (
                f"عدد صور الأسئلة داخل PDF ({extracted_image_count}) لا يطابق "
                f"المرفقات الصالحة ({valid_attachments})."
            )
        )
        text_layer_message = (
            "يحتوي PDF على طبقة نصية قابلة للاستخراج."
            if text_layer_passed
            else "لم تُكتشف طبقة نصية قابلة للاستخراج داخل PDF."
        )

    checks.extend(
        [
            FullExamExportCheck(
                code="pdf_pages",
                passed=page_count > 0,
                message=(
                    f"يحتوي PDF على {page_count} صفحة قابلة للقراءة."
                    if page_count > 0
                    else "لا يحتوي PDF على صفحات قابلة للقراءة."
                ),
            ),
            FullExamExportCheck(
                code="pdf_question_headings",
                passed=heading_passed,
                message=heading_message,
            ),
            FullExamExportCheck(
                code="pdf_attachments_once",
                passed=attachment_passed,
                message=attachment_message,
            ),
            FullExamExportCheck(
                code="pdf_visual_question_coverage",
                passed=not missing_visual_questions,
                message=(
                    "كل سؤال يعتمد على رسم أو مخطط يملك مرفقًا بصريًا صالحًا في PDF."
                    if not missing_visual_questions
                    else (
                        "توجد أسئلة تعتمد على رسم أو مخطط دون مرفق بصري صالح في PDF: "
                        + "، ".join(missing_visual_questions)
                    )
                ),
            ),
            FullExamExportCheck(
                code="pdf_text_layer",
                passed=text_layer_passed,
                message=text_layer_message,
            ),
        ]
    )

    invalid_attachments = int(expected["attachments"]) - valid_attachments
    if invalid_attachments:
        warnings.append(
            f"تعذر إدراج {invalid_attachments} مرفق بسبب بيانات صورة غير صالحة."
        )

    status = (
        FullExamExportArtifactStatus.accepted
        if all(check.passed for check in checks) and not warnings
        else FullExamExportArtifactStatus.needs_review
    )
    if not signature_valid or manifest is None or page_count == 0:
        status = FullExamExportArtifactStatus.failed

    return FullExamExportFormatSummary(
        format=ExportFormat.pdf,
        status=status,
        byte_size=len(artifact_bytes),
        page_count=page_count or None,
        exported_question_count=exported_question_count,
        exported_part_count=exported_part_count,
        exported_attachment_count=exported_attachment_count,
        detected_total_marks=detected_total_marks,
        question_order=question_order,
        checks=checks,
        warnings=warnings,
    )


def _unique_formats(formats: Iterable[ExportFormat]) -> list[ExportFormat]:
    result: list[ExportFormat] = []
    for item in formats:
        if item not in result:
            result.append(item)
    return result


def build_full_exam_export_report(
    project: ProjectSession,
    artifact_format: ExportFormat,
    artifact_bytes: bytes,
    existing_report: FullExamExportReport | None = None,
) -> FullExamExportReport:
    """Inspect one real artifact and merge it into the persisted exam report."""

    format_summary = (
        _inspect_docx(project, artifact_bytes)
        if artifact_format == ExportFormat.docx
        else _inspect_pdf(project, artifact_bytes)
    )

    summaries_by_format = {
        summary.format: summary
        for summary in (existing_report.formats if existing_report else [])
    }
    summaries_by_format[artifact_format] = format_summary

    requested_formats = _unique_formats(
        project.metadata.export_formats or [ExportFormat.docx, ExportFormat.pdf]
    )
    summaries = [
        summaries_by_format[item]
        for item in (ExportFormat.docx, ExportFormat.pdf)
        if item in summaries_by_format
    ]
    generated_formats = [summary.format for summary in summaries]
    accepted_formats = [
        summary.format
        for summary in summaries
        if summary.status == FullExamExportArtifactStatus.accepted
    ]
    needs_review_formats = [
        summary.format
        for summary in summaries
        if summary.status == FullExamExportArtifactStatus.needs_review
    ]
    failed_formats = [
        summary.format
        for summary in summaries
        if summary.status == FullExamExportArtifactStatus.failed
    ]

    questions = _active_questions(project)
    expected = _expected_structure(project)
    reported_marks = parse_declared_total_marks(project.metadata.total_marks)
    missing_formats = [
        item
        for item in requested_formats
        if item not in generated_formats
    ]

    checks = [
        FullExamExportCheck(
            code="active_questions",
            passed=bool(questions),
            message=(
                f"توجد {len(questions)} أسئلة نشطة قابلة للتصدير."
                if questions
                else "لا توجد أسئلة نشطة قابلة للتصدير."
            ),
        ),
        FullExamExportCheck(
            code="question_order",
            passed=(
                len({question.order_index for question in questions})
                == len(questions)
                and [question.order_index for question in questions]
                == sorted(question.order_index for question in questions)
            ),
            message="ترتيب الأسئلة النشطة فريد ومستقر قبل التصدير.",
        ),
        FullExamExportCheck(
            code="paper_total_matches",
            passed=marks_policy_resolves_total(
                project.metadata,
                int(expected["marks"]),
            ),
            message=(
                "سياسة درجة الورقة محسومة ومتوافقة مع التصدير."
                if marks_policy_resolves_total(
                    project.metadata,
                    int(expected["marks"]),
                )
                else (
                    f"مجموع الأسئلة ({expected['marks']}) لا يطابق الدرجة "
                    f"المعلنة ({reported_marks}) ولم تُحدَّد سياسة للحسم."
                )
            ),
        ),
        FullExamExportCheck(
            code="requested_formats_generated",
            passed=not missing_formats,
            message=(
                "تم إنشاء جميع صيغ التصدير المطلوبة."
                if not missing_formats
                else "لم تُنشأ بعد الصيغ المطلوبة: "
                + "، ".join(item.value.upper() for item in missing_formats)
            ),
        ),
        FullExamExportCheck(
            code="artifacts_accepted",
            passed=not failed_formats and not needs_review_formats,
            message=(
                "اجتازت ملفات التصدير المنشأة فحوص البنية."
                if not failed_formats and not needs_review_formats
                else "توجد ملفات تصدير تحتاج مراجعة أو فشلت فحوصها."
            ),
        ),
    ]

    warnings: list[str] = []
    if project.full_exam_intake_report is None:
        warnings.append("لا يوجد تقرير قبول لبنية الورقة الكاملة.")
    elif project.full_exam_intake_report.status != FullExamIntakeStatus.accepted:
        warnings.append("تقرير إدخال الورقة الكاملة لم يصل إلى حالة accepted.")

    if project.full_exam_translation_report is None:
        warnings.append("لا يوجد تقرير قبول لترجمة الورقة الكاملة.")
    elif (
        project.full_exam_translation_report.status
        != FullExamTranslationAcceptanceStatus.accepted
    ):
        warnings.append("تقرير ترجمة الورقة الكاملة لم يصل إلى حالة accepted.")

    for summary in summaries:
        warnings.extend(
            f"{summary.format.value.upper()}: {warning}"
            for warning in summary.warnings
        )

    if failed_formats or not questions:
        status = FullExamExportAcceptanceStatus.failed
    elif missing_formats:
        status = FullExamExportAcceptanceStatus.incomplete
    elif needs_review_formats or warnings or not all(check.passed for check in checks):
        status = FullExamExportAcceptanceStatus.needs_review
    else:
        status = FullExamExportAcceptanceStatus.accepted

    return FullExamExportReport(
        status=status,
        requested_formats=requested_formats,
        generated_formats=generated_formats,
        accepted_formats=accepted_formats,
        needs_review_formats=needs_review_formats,
        failed_formats=failed_formats,
        active_question_count=len(questions),
        expected_total_marks=int(expected["marks"]),
        expected_part_count=int(expected["parts"]),
        expected_attachment_count=int(expected["attachments"]),
        source_page_linked_questions=sum(
            bool(question.source_page_numbers)
            for question in questions
        ),
        multi_page_questions=sum(
            len(question.source_page_numbers) > 1
            for question in questions
        ),
        formats=summaries,
        checks=checks,
        warnings=warnings,
    )
