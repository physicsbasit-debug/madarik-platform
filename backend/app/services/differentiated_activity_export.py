from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.models.differentiated_activity import (
    DifferentiatedActivity,
    DifferentiatedActivityExportResponse,
    DifferentiatedActivityPreview,
    DifferentiationLevel,
)

EXPORT_DIR = Path("data/exports/differentiated-activities")

LEVEL_LABELS = {
    DifferentiationLevel.support: "دعم",
    DifferentiationLevel.core: "أساسي",
    DifferentiationLevel.extension: "إثراء",
}


def build_activity_preview(
    activity: DifferentiatedActivity,
) -> DifferentiatedActivityPreview:
    issues: list[str] = []
    if not activity.title.strip():
        issues.append("عنوان النشاط غير موجود.")
    if not activity.objective.strip():
        issues.append("هدف النشاط غير موجود.")
    if not activity.instructions.strip():
        issues.append("تعليمات النشاط غير موجودة.")
    if not activity.success_criteria:
        issues.append("معايير النجاح غير موجودة.")

    return DifferentiatedActivityPreview(
        id=activity.id,
        title=activity.title,
        level=activity.level,
        level_label=LEVEL_LABELS[activity.level],
        grade=activity.grade,
        science_domain=activity.science_domain,
        subject_id=activity.subject_id,
        unit_id=activity.unit_id,
        lesson_id=activity.lesson_id,
        objective=activity.objective,
        instructions=activity.instructions,
        success_criteria=activity.success_criteria,
        estimated_minutes=activity.estimated_minutes,
        materials=activity.materials,
        export_ready=not issues,
        issues=issues,
    )


def _set_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def _heading(document: Document, text: str, size: int = 14) -> None:
    paragraph = document.add_paragraph()
    _set_rtl(paragraph)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _build_docx(preview: DifferentiatedActivityPreview, path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    title = document.add_paragraph()
    _set_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(preview.title)
    run.bold = True
    run.font.size = Pt(18)

    meta = document.add_table(rows=2, cols=3)
    meta.style = "Table Grid"
    values = (
        ("المستوى", preview.level_label),
        ("الصف", str(preview.grade)),
        ("الزمن", f"{preview.estimated_minutes} دقيقة"),
    )
    for index, (label, value) in enumerate(values):
        meta.cell(0, index).text = label
        meta.cell(1, index).text = value
        for row in (0, 1):
            for paragraph in meta.cell(row, index).paragraphs:
                _set_rtl(paragraph)

    _heading(document, "الهدف")
    paragraph = document.add_paragraph(preview.objective)
    _set_rtl(paragraph)

    _heading(document, "التعليمات")
    paragraph = document.add_paragraph(preview.instructions)
    _set_rtl(paragraph)

    _heading(document, "معايير النجاح")
    for criterion in preview.success_criteria:
        paragraph = document.add_paragraph(criterion, style="List Bullet")
        _set_rtl(paragraph)

    _heading(document, "الأدوات والمواد")
    if preview.materials:
        for material in preview.materials:
            paragraph = document.add_paragraph(material, style="List Bullet")
            _set_rtl(paragraph)
    else:
        paragraph = document.add_paragraph("لا توجد أدوات محددة.")
        _set_rtl(paragraph)

    _heading(document, "مساحة عمل الطالب")
    for _ in range(8):
        paragraph = document.add_paragraph(
            "................................................................................"
        )
        _set_rtl(paragraph)

    document.save(path)


def _register_font() -> str:
    candidates = (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
    )
    for candidate in candidates:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont("ActivityArabic", str(candidate)))
            return "ActivityArabic"
    return "Helvetica"


def _build_pdf(preview: DifferentiatedActivityPreview, path: Path) -> None:
    font_name = _register_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ActivityTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=18,
        leading=24,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "ActivityHeading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=14,
        leading=20,
        alignment=TA_RIGHT,
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ActivityBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=11,
        leading=18,
        alignment=TA_RIGHT,
    )

    story = [
        Paragraph(preview.title, title_style),
        Table(
            [
                ["المستوى", "الصف", "الزمن"],
                [
                    preview.level_label,
                    str(preview.grade),
                    f"{preview.estimated_minutes} دقيقة",
                ],
            ],
            colWidths=[5.5 * cm, 5.5 * cm, 5.5 * cm],
            style=TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.5, "black"),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                ]
            ),
        ),
        Spacer(1, 0.4 * cm),
        Paragraph("الهدف", heading_style),
        Paragraph(preview.objective, body_style),
        Paragraph("التعليمات", heading_style),
        Paragraph(preview.instructions, body_style),
        Paragraph("معايير النجاح", heading_style),
    ]

    for criterion in preview.success_criteria:
        story.append(Paragraph(f"• {criterion}", body_style))

    story.append(Paragraph("الأدوات والمواد", heading_style))
    if preview.materials:
        for material in preview.materials:
            story.append(Paragraph(f"• {material}", body_style))
    else:
        story.append(Paragraph("لا توجد أدوات محددة.", body_style))

    story.append(Paragraph("مساحة عمل الطالب", heading_style))
    for _ in range(8):
        story.append(Paragraph("." * 95, body_style))

    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )
    document.build(story)


def export_activity(
    activity: DifferentiatedActivity,
    output_format: str,
) -> DifferentiatedActivityExportResponse:
    preview = build_activity_preview(activity)
    if not preview.export_ready:
        return DifferentiatedActivityExportResponse(
            activity_id=activity.id,
            format=output_format,
            filename="",
            path="",
            export_ready=False,
            issues=preview.issues,
        )

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    safe_title = activity.title.strip().replace("/", "-") or "activity"
    filename = f"{safe_title}-{activity.id}.{output_format}"
    path = EXPORT_DIR / filename

    if output_format == "docx":
        _build_docx(preview, path)
    else:
        _build_pdf(preview, path)

    return DifferentiatedActivityExportResponse(
        activity_id=activity.id,
        format=output_format,
        filename=filename,
        path=str(path),
        export_ready=True,
        issues=[],
    )
